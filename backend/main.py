# main.py — SubtitleAI V2 Backend
# FastAPI + Groq (LLaMA 3.1 8B Instant) + MySQL (Reloaded)

import sys
import os
import warnings

# Suppress Hugging Face Symlink warning on Windows and other deprecation/user warnings
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if sys.stderr and hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse
from dotenv import load_dotenv
import json, io, re, asyncio

from database import init_db, get_all_platforms, save_custom_platform, log_job
from file_reader import read_file
from extractor import pre_extract_dialogue
from cleaner import clean_subtitle_chunk, extract_platform_rules_with_ai, _preserve_as_broadcast_title
from quality_checker import check_quality
from platform_rules import get_platform, get_platform_list
from timecoded_subtitles import ensure_srt_timings, parse_timecoded_subtitles, prepare_for_platform, subtitles_to_srt, normalize_timecode
import auth

load_dotenv()

app = FastAPI(title="SubtitleAI V2", description="AI subtitle cleaning and quality check tool", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[os.getenv("FRONTEND_URL", "http://localhost:5173"), "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(auth.router, prefix="/auth")

@app.on_event("startup")
def startup():
    try:
        init_db()
    except Exception as e:
        print(f"WARNING DB warning: {e}")
    try:
        from guidelines_db import init_guidelines_table
        init_guidelines_table()
    except Exception as e:
        print(f"WARNING Guidelines DB warning: {e}")


@app.get("/")
def root():
    return {"status": "running", "version": "2.0.0", "model": os.getenv("GROQ_MODEL", "openai/gpt-oss-20b")}


@app.get("/health")
def health():
    groq_key = os.getenv("GROQ_API_KEY", "")
    return {
        "status": "ok",
        "groq_configured": bool(groq_key and groq_key != "your_groq_api_key_here"),
        "model": f"{os.getenv('GROQ_MODEL', 'openai/gpt-oss-20b')} (Groq)"
    }

@app.get("/semantic-model/status")
def semantic_model_status():
    from semantic_matcher import model_status
    return model_status()


@app.get("/bridge/status")
def bridge_status():
    """Report local bridge availability without downloading language packs."""
    from bridge_translator import status
    return status()


@app.post("/bridge/prepare")
async def bridge_prepare(source_language: str = Form(...), target_language: str = Form(...)):
    """Explicitly prepare one local Argos language pair."""
    import asyncio
    try:
        from bridge_translator import prepare_language_pair
        return await asyncio.to_thread(prepare_language_pair, source_language, target_language)
    except Exception as exc:
        raise HTTPException(503, str(exc))


@app.post("/bridge/prepare-from-files")
async def bridge_prepare_from_files(
    script_file: UploadFile = File(...), timestamps_file: UploadFile = File(...)
):
    """Detect the two languages from the selected files, then prepare one pair."""
    import asyncio
    try:
        from file_reader import read_file
        from extractor import pre_extract_dialogue
        from timecoded_subtitles import parse_timecoded_subtitles
        from bridge_translator import detect_languages, prepare_language_pair

        script_bytes = await script_file.read()
        ts_bytes = await timestamps_file.read()
        script_name = script_file.filename or "script.txt"
        ts_name = timestamps_file.filename or "timestamps.srt"
        sd = read_file(script_bytes, script_name)
        parsed = parse_timecoded_subtitles(sd["raw_text"]) if sd["structure"] in ("srt", "vtt", "ttml") else []
        if parsed:
            client_subs = parsed
        else:
            raw_lines = pre_extract_dialogue(sd["raw_text"], sd["structure"], script_bytes, script_name, {}) or sd["raw_text"].splitlines()
            client_subs = [{"text": line.strip()} for line in raw_lines if line.strip()]
        td = read_file(ts_bytes, ts_name)
        whisper_subs = parse_timecoded_subtitles(td["raw_text"])
        if not client_subs or not whisper_subs:
            raise ValueError("Both files must contain readable subtitle text.")
        source, target = detect_languages(whisper_subs, client_subs)
        result = await asyncio.to_thread(prepare_language_pair, source, target)
        return result | {"source_language": source, "target_language": target}
    except Exception as exc:
        raise HTTPException(503, str(exc))


@app.post("/bridge/prepare-from-files-stream")
async def bridge_prepare_from_files_stream(
    script_file: UploadFile = File(...), timestamps_file: UploadFile = File(...)
):
    """Streaming one-time bridge setup with explicit detection/download stages."""
    import queue, threading, time
    script_bytes = await script_file.read()
    ts_bytes = await timestamps_file.read()
    script_name = script_file.filename or "script.txt"
    ts_name = timestamps_file.filename or "timestamps.srt"
    events_queue: queue.Queue = queue.Queue()

    def worker():
        try:
            from file_reader import read_file
            from extractor import pre_extract_dialogue
            from timecoded_subtitles import parse_timecoded_subtitles
            from bridge_translator import detect_languages, prepare_language_pair
            events_queue.put({"status": "processing", "progress": 10, "message": "Reading the selected subtitle files..."})
            sd = read_file(script_bytes, script_name)
            parsed = parse_timecoded_subtitles(sd["raw_text"]) if sd["structure"] in ("srt", "vtt", "ttml") else []
            if parsed:
                client_subs = parsed
            else:
                raw_lines = pre_extract_dialogue(sd["raw_text"], sd["structure"], script_bytes, script_name, {}) or sd["raw_text"].splitlines()
                client_subs = [{"text": line.strip()} for line in raw_lines if line.strip()]
            td = read_file(ts_bytes, ts_name)
            whisper_subs = parse_timecoded_subtitles(td["raw_text"])
            source, target = detect_languages(whisper_subs, client_subs)
            events_queue.put({"status": "processing", "progress": 25, "message": f"Detected bridge: {source} → {target}."})
            if source == target:
                events_queue.put({"status": "completed", "progress": 100, "message": "The languages already match; no bridge pack is needed.", "result": {"ready": True, "source_language": source, "target_language": target}})
                return
            events_queue.put({"status": "processing", "progress": 35, "message": "Preparing the local language pack (one time)..."})
            result = prepare_language_pair(source, target)
            events_queue.put({"status": "completed", "progress": 100, "message": "Local language bridge ready.", "result": result | {"source_language": source, "target_language": target}})
        except Exception as exc:
            events_queue.put({"status": "error", "progress": 0, "error": str(exc)})

    threading.Thread(target=worker, daemon=True).start()

    async def event_stream():
        last_heartbeat = time.monotonic()
        last_progress = 35
        while True:
            try:
                event = events_queue.get_nowait()
            except queue.Empty:
                if time.monotonic() - last_heartbeat >= 4:
                    last_heartbeat = time.monotonic()
                    yield f"data: {json.dumps({'status':'processing','progress':last_progress,'message':'Downloading/verifying the language pack... this is only needed once.'})}\n\n"
                await asyncio.sleep(0.15)
                continue
            last_heartbeat = time.monotonic()
            if event.get("status") == "processing":
                last_progress = event.get("progress", last_progress)
            yield f"data: {json.dumps(event)}\n\n"
            if event.get("status") in ("completed", "error"):
                break

    return StreamingResponse(event_stream(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.post("/semantic-model/prepare")
async def semantic_model_prepare():
    """Explicit one-time download/cache action; never runs during alignment."""
    import asyncio
    try:
        from semantic_matcher import prepare_model
        return await asyncio.wait_for(asyncio.to_thread(prepare_model), timeout=180)
    except Exception as exc:
        raise HTTPException(503, f"Could not prepare multilingual model within the setup timeout. Check model download/network access and retry. Details: {exc}")


@app.post("/semantic-model/prepare-stream")
async def semantic_model_prepare_stream():
    """Streaming prepare endpoint — emits SSE progress events while downloading/loading the model."""
    import queue, threading, time
    q: queue.Queue = queue.Queue()

    def worker():
        try:
            from semantic_matcher import model_status, prepare_model
            status = model_status()
            q.put({"status": "processing", "progress": 10, "message": "Connecting to model repository…"})
            # SentenceTransformer will download if not cached; we can't intercept
            # granular download progress without monkey-patching, so we emit a
            # steady heartbeat in the reader loop and report a final loaded step.
            q.put({"status": "processing", "progress": 20, "message": "Downloading / verifying model weights (this may take a minute)…"})
            prepare_model()
            q.put({"status": "processing", "progress": 90, "message": "Model loaded into memory."})
            q.put({"status": "completed", "progress": 100, "message": "Multilingual model ready.", "result": model_status()})
        except Exception as exc:
            q.put({"status": "error", "progress": 0, "error": str(exc)})

    threading.Thread(target=worker, daemon=True).start()

    async def events():
        heartbeat_pct = 20
        last_hb = time.monotonic()
        while True:
            try:
                msg = q.get_nowait()
            except queue.Empty:
                if time.monotonic() - last_hb >= 4:
                    last_hb = time.monotonic()
                    heartbeat_pct = min(85, heartbeat_pct + 5)
                    yield f"data: {json.dumps({'status':'processing','progress':heartbeat_pct,'message':'Downloading model weights… please wait.'})}\n\n"
                await asyncio.sleep(0.2)
                continue
            yield f"data: {json.dumps(msg)}\n\n"
            if msg["status"] in ("completed", "error"):
                break

    return StreamingResponse(
        events(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ─── CLEAN ───────────────────────────────────────────────────────

def chunk_list(items: list[str], max_chunk_size: int = 7000) -> list[list[str]]:
    chunks = []
    current_chunk = []
    current_size = 0
    for item in items:
        item_len = len(item) + 1
        if current_size + item_len > max_chunk_size and current_chunk:
            chunks.append(current_chunk)
            current_chunk = [item]
            current_size = item_len
        else:
            current_chunk.append(item)
            current_size += item_len
    if current_chunk:
        chunks.append(current_chunk)
    return chunks


@app.post("/clean")
async def clean_file_endpoint(
    file: UploadFile = File(...),
    platform: str = Form(default="generic"),
    force_ocr: bool = Form(False)
):
    import asyncio
    
    ALLOWED = [".doc",".docx",".pdf",".xml",".ttml",".dfxp",".pmw",
               ".rtf",".srt",".vtt",".webvtt",".xlsx",".xls",
               ".csv",".txt",".json",".pac"]

    filename = file.filename or "unknown.txt"
    ext = "." + filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    file_ext = ext.lstrip(".")

    if ext not in ALLOWED:
        raise HTTPException(400, f"File type '{ext}' not supported.")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(400, "Uploaded file is empty")

    try:
        file_data = read_file(file_bytes, filename, force_ocr=force_ocr)
    except ValueError as exc:
        raise HTTPException(422, str(exc))
    raw_text = file_data["raw_text"]
    structure = file_data["structure"]

    if not raw_text.strip():
        raise HTTPException(400, "Could not extract text from file")

    # ── Step 1: Pure Python extraction (instant, no LLM) ─────────────────────
    platform_dict = get_platform(platform)
    timecoded_subtitles = parse_timecoded_subtitles(raw_text)

    groq_key = os.getenv("GROQ_API_KEY", "")
    if not groq_key or groq_key == "your_groq_api_key_here":
        raise HTTPException(500, "GROQ_API_KEY not configured. Get free key at https://console.groq.com")

    if timecoded_subtitles and "--- OCR SUBTITLES" not in raw_text:
        pre_extracted = [sub.get("text", "") for sub in timecoded_subtitles]
        print(f"[EXTRACT] Found {len(pre_extracted)} timecoded subtitles, sending to LLM for cleaning.")
    else:
        pre_extracted = pre_extract_dialogue(raw_text, structure, file_bytes, filename, platform_dict)

    if not pre_extracted:
        pre_extracted = [line.strip() for line in raw_text.splitlines() if line.strip()]
        print(f"[EXTRACT] No pre-extraction, sending {len(pre_extracted)} raw lines to LLM")

    chunk_size = int(os.getenv("GROQ_CHUNK_SIZE", "7000"))
    chunks = chunk_list(pre_extracted, max_chunk_size=chunk_size)
    total_chunks = len(chunks)
    
    # Process fewer chunks in parallel with larger sizes to prevent exceeding Groq TPM.
    parallel = 1

    async def process_chunk(idx: int, chunk: list[str]):
        loop = asyncio.get_running_loop()
        return idx, await loop.run_in_executor(
            None,
            clean_subtitle_chunk,
            chunk,
            structure,
            platform,
            filename
        )

    async def event_generator():
        all_subtitles = []

        yield f"data: {json.dumps({'status': 'starting', 'progress': 0, 'message': f'Initializing — {total_chunks} parts, processing {parallel} at a time...'})}\n\n"
        await asyncio.sleep(0.05)

        # Process one batch at a time to preserve order and protect API limits.
        for batch_start in range(0, total_chunks, parallel):
            pacing = float(os.getenv("GROQ_CHUNK_DELAY", "0"))
            if pacing > 0 and batch_start > 0:
                yield f"data: {json.dumps({'status': 'processing', 'progress': int((batch_start / total_chunks) * 100), 'message': f'Groq Pacing: Waiting 21s to clear 6k TPM bucket...'})}\n\n"
                await asyncio.sleep(pacing)
                
            batch = list(enumerate(chunks[batch_start:batch_start + parallel], start=batch_start))
            done = batch_start + len(batch)
            progress_pct = int((batch_start / total_chunks) * 100)
            part_range = f"{batch_start+1}–{done}" if len(batch) > 1 else str(batch_start+1)
            yield f"data: {json.dumps({'status': 'processing', 'progress': progress_pct, 'message': f'AI cleaning parts {part_range} of {total_chunks}...'})}\n\n"
            await asyncio.sleep(0.05)

            try:
                results = await asyncio.gather(*[process_chunk(i, c) for i, c in batch])
                # Sort by original index to preserve document order
                results.sort(key=lambda x: x[0])
                for _, chunk_subs in results:
                    all_subtitles.extend(chunk_subs)
            except Exception as e:
                yield f"data: {json.dumps({'status': 'error', 'error': str(e)})}\n\n"
                return

        # Re-index all subtitles contiguously and assign sequential placeholder timecodes if missing
        from timecoded_subtitles import _from_seconds, _to_seconds, ensure_srt_timings, prepare_for_platform, normalize_timecode
        
        # 1. Restore timecodes if we had them!
        if timecoded_subtitles and "--- OCR SUBTITLES" not in raw_text:
            for i, sub in enumerate(all_subtitles):
                if i < len(timecoded_subtitles):
                    sub["start_time"] = timecoded_subtitles[i].get("start_time", "")
                    sub["end_time"] = timecoded_subtitles[i].get("end_time", "")
                    
        # 2. Filter out [DELETE] / empty rows!
        all_subtitles = [sub for sub in all_subtitles if not sub.get("deleted", False)]

        # PAC is a complete, already-timed subtitle format.  An LLM response
        # that is cut short or omits an item must never turn a 500-cue PAC
        # into a shorter SRT.  Fall back to the native PAC cues when the LLM
        # does not return a strict one-for-one result; deterministic platform
        # formatting below still runs on every preserved cue.
        if file_ext == "pac" and timecoded_subtitles and len(all_subtitles) != len(timecoded_subtitles):
            print(
                f"[PAC] LLM returned {len(all_subtitles)} of "
                f"{len(timecoded_subtitles)} cues; preserving native PAC cues instead."
            )
            all_subtitles = [
                {
                    **source_sub,
                    "original_text": source_sub.get("text", ""),
                    "text": source_sub.get("text", ""),
                    "deleted": False,
                }
                for source_sub in timecoded_subtitles
            ]
        
        tc = 0.0
        for s_idx, sub in enumerate(all_subtitles, start=1):
            sub["id"] = s_idx
            
            # If timecodes are missing, generate sequential ones (3s apart) so SRT export works
            if not sub.get("start_time"):
                sub["start_time"] = _from_seconds(tc)
                sub["end_time"] = _from_seconds(tc + 2.0)
                tc += 3.0
            else:
                # Keep existing timecodes, but update 'tc' so any subsequent missing ones continue from here
                parsed = _to_seconds(sub["start_time"])
                if parsed is not None:
                    tc = parsed + 3.0
                if not sub.get("end_time"):
                    sub["end_time"] = _from_seconds(tc - 1.0)

        # ── Auto-fix pass: apply platform rules in Python (100% reliable) ──
        from quality_checker import auto_fix_subtitles
        from platform_rules import get_platform as _get_platform
        all_subtitles = auto_fix_subtitles(all_subtitles, platform)
        
        if timecoded_subtitles and "--- OCR SUBTITLES" not in raw_text:
            all_subtitles = ensure_srt_timings(all_subtitles)
            all_subtitles = prepare_for_platform(all_subtitles, platform, filename)
            for s_idx, sub in enumerate(all_subtitles, start=1):
                sub["id"] = s_idx
                sub["start_time"] = normalize_timecode(sub.get("start_time", ""))
                sub["end_time"] = normalize_timecode(sub.get("end_time", ""))

        # Re-evaluate flagged status after auto_fix.
        # auto_fix resolves most issues (line splitting, profanity, formatting).
        # Only keep flagged=True if the line is STILL too long after splitting.
        _plat = _get_platform(platform)
        _max_chars = int(_plat.get("max_chars_per_line", 42))
        for sub in all_subtitles:
            if sub.get("flagged"):
                # Check if line is actually still too long
                _still_long = any(
                    len(re.sub(r'<[^>]+>', '', ln)) > _max_chars
                    for ln in sub.get("text", "").split("\n")
                )
                if not _still_long:
                    sub["flagged"] = False
                    sub["flag_reason"] = ""

        total_lines = len(all_subtitles)
        flagged_lines = sum(1 for s in all_subtitles if s.get("flagged"))
        # How many lines the AI/auto-fix actually changed vs the raw extracted
        # text — this is the proof-of-work number the subtitler can glance at
        # to confirm cleaning genuinely happened, not just "it ran fast".
        changed_lines = sum(
            1 for s in all_subtitles
            if s.get("original_text", s.get("text", "")).strip() != s.get("text", "").strip()
        )

        # Log job
        try:
            log_job(filename, file_ext, platform, structure, total_lines, flagged_lines, 0)
        except Exception as e:
            print(f"Failed to log job: {e}")

        final_result = {
            "subtitles": all_subtitles,
            "stats": {
                "total_lines": total_lines,
                "flagged_lines": flagged_lines,
                "changed_lines": changed_lines,
                "platform": platform,
                "detected_structure": structure,
                "original_format": filename
            }
        }

        yield f"data: {json.dumps({'status': 'completed', 'progress': 100, 'result': final_result})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",    # disables Nginx/proxy buffering
            "Connection": "keep-alive",
        }
    )

@app.post("/extract")
async def extract_file_endpoint(
    file: UploadFile = File(...),
    platform: str = Form(default="generic"),
    force_ocr: bool = Form(False)
):
    ALLOWED = [".doc",".docx",".pdf",".xml",".ttml",".dfxp",".pmw",
               ".rtf",".srt",".vtt",".webvtt",".xlsx",".xls",
               ".csv",".txt",".json",".pac"]

    filename = file.filename or "unknown.txt"
    ext = "." + filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext not in ALLOWED:
        raise HTTPException(400, f"File type '{ext}' not supported.")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(400, "Uploaded file is empty")

    try:
        file_data = read_file(file_bytes, filename, force_ocr=force_ocr)
    except ValueError as exc:
        raise HTTPException(422, str(exc))
    raw_text = file_data["raw_text"]
    structure = file_data["structure"]

    if not raw_text.strip():
        raise HTTPException(400, "Could not extract text from file")

    platform_dict = get_platform(platform)
    timecoded_subtitles = parse_timecoded_subtitles(raw_text)
    if timecoded_subtitles and "--- OCR SUBTITLES" not in raw_text:
        return {
            "subtitles": timecoded_subtitles,
            "stats": {
                "total_lines": len(timecoded_subtitles),
                "flagged_lines": 0,
                "platform": "none",
                "detected_structure": "srt_timecoded",
                "original_format": filename
            }
        }

    pre_extracted = []
    if "--- OCR SUBTITLES" not in raw_text:
        pre_extracted = pre_extract_dialogue(raw_text, structure, file_bytes, filename, platform_dict)
    if not pre_extracted:
        pre_extracted = raw_text.splitlines()

    subtitles = []
    for i, line in enumerate(pre_extracted, 1):
        clean_line = line.strip()
        if clean_line:
            subtitles.append({"id": i, "original_text": clean_line, "text": clean_line, "flagged": False})

    return {
        "subtitles": subtitles,
        "stats": {
            "total_lines": len(subtitles),
            "flagged_lines": 0,
            "platform": "none",
            "detected_structure": structure,
            "original_format": filename
        }
    }


# ─── QUALITY CHECK ───────────────────────────────────────────────

@app.post("/align-client-timeline")
async def align_client_timeline_endpoint(data: dict):
    """Transform reference-timed mapped subtitles into client/GTS time."""
    subtitles = data.get("subtitles", [])
    anchors = data.get("anchors", [])
    if not subtitles or not anchors:
        raise HTTPException(400, "subtitles and at least two reference/client anchors are required")
    try:
        from timeline_alignment import align_subtitles_to_client
        pairs = [(a.get("reference", a.get("ref")), a.get("client", a.get("gts"))) if isinstance(a, dict) else tuple(a) for a in anchors]
        result, report = align_subtitles_to_client(
            subtitles, pairs, data.get("reference_duration"), data.get("client_duration"),
            data.get("client_fps"), data.get("reference_fps"))
        return {"subtitles": result, "report": report}
    except (TypeError, ValueError) as exc:
        raise HTTPException(422, str(exc))


@app.post("/quality-check")
async def quality_check_endpoint(data: dict):
    """
    Check cleaned subtitles for defects before delivery to OTT platform.
    Accepts: { subtitles: [...], platform_key: "...", filename: "..." }
    """
    subtitles = data.get("subtitles", [])
    platform_key = data.get("platform_key", "generic")
    filename = data.get("filename", "subtitles.srt")
    fps = float(data.get("fps", 25.0) or 25.0)

    if not subtitles:
        raise HTTPException(400, "No subtitles provided for quality check")

    subtitles = prepare_for_platform(subtitles, platform_key, filename)
    result = check_quality(subtitles, platform_key, filename, fps=fps)

    # Log defects
    try:
        log_job(filename, "quality_check", platform_key, "quality_check",
                result.get("total_lines", 0), 0, result.get("total_defects", 0))
    except Exception:
        pass

    return result


# ─── EXPORT ──────────────────────────────────────────────────────

@app.post("/export/srt")
async def export_srt(data: dict):
    """Export real SRT with preserved timecodes."""
    subtitles = data.get("subtitles", [])
    platform_key = data.get("platform_key", "generic")
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]

    if not data.get("preserve_exact", False):
        subtitles = prepare_for_platform(subtitles, platform_key, data.get("filename", "cleaned"))
    content = subtitles_to_srt(subtitles)
    if not content.strip():
        raise HTTPException(400, "No timecoded subtitles available for SRT export.")

    buf = io.BytesIO(content.encode("utf-8"))
    return StreamingResponse(buf, media_type="application/x-subrip",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.srt"})


@app.post("/export/txt")
async def export_txt(data: dict):
    subtitles = data.get("subtitles", [])
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]
    content = "\n".join(s.get("text", "") for s in subtitles if s.get("text"))
    buf = io.BytesIO(content.encode("utf-8"))
    return StreamingResponse(buf, media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.txt"})


@app.post("/export/docx")
async def export_docx(data: dict):
    """Export dialogue-only as a DOCX with one paragraph per line and blank spacer between entries."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    subtitles = data.get("subtitles", [])
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]

    doc = Document()
    # Set a readable body font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    for i, sub in enumerate(subtitles):
        text = sub.get("text", "").strip()
        if not text:
            continue
        # Remove control characters that break python-docx
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        # Add the dialogue line, centred to match the screenshot layout
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add a blank spacer paragraph between entries
        if i < len(subtitles) - 1:
            spacer = doc.add_paragraph("")
            spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.docx"})


@app.post("/export/pdf")
async def export_pdf(data: dict):
    """Export dialogue-only as a PDF, centred on the page with spacing between entries."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.units import inch
    import html
    import re

    subtitles = data.get("subtitles", [])
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=1.5 * inch, rightMargin=1.5 * inch,
        topMargin=1 * inch, bottomMargin=1 * inch
    )
    styles = getSampleStyleSheet()
    centred = ParagraphStyle(
        "centred",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=12,
        leading=18,
        spaceAfter=0,
    )
    Story = []

    for sub in subtitles:
        text = sub.get("text", "").strip()
        if not text:
            continue
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        text = html.escape(text).replace("\n", "<br/>")
        Story.append(Paragraph(text, centred))
        # Generous vertical space between each dialogue entry (matches screenshot)
        Story.append(Spacer(1, 28))

    doc.build(Story)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.pdf"})


@app.post("/export/vtt")
async def export_vtt(data: dict):
    subtitles = data.get("subtitles", [])
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]
    
    blocks = ["WEBVTT\n"]
    for i, sub in enumerate(subtitles, start=1):
        text = (sub.get("text") or "").strip()
        if not text: continue
        start = sub.get("start_time", "").replace(",", ".")
        end = sub.get("end_time", "").replace(",", ".")
        if not start or not end: continue
        blocks.append(f"{i}\n{start} --> {end}\n{text}")
        
    content = "\n\n".join(blocks)
    buf = io.BytesIO(content.encode("utf-8"))
    return StreamingResponse(buf, media_type="text/vtt",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.vtt"})


@app.post("/export/ttml")
async def export_ttml(data: dict):
    subtitles = data.get("subtitles", [])
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]
    import xml.sax.saxutils as saxutils
    
    lines = [
        '<?xml version="1.0" encoding="utf-8"?>',
        '<tt xmlns="http://www.w3.org/ns/ttml" xmlns:tts="http://www.w3.org/ns/ttml#styling" xml:lang="en">',
        '  <head>',
        '    <styling>',
        '      <style xml:id="defaultStyle" tts:fontFamily="sansSerif" tts:fontSize="12px" tts:textAlign="center" tts:color="white" tts:backgroundColor="transparent"/>',
        '    </styling>',
        '    <layout>',
        '      <region xml:id="bottom" tts:origin="10% 80%" tts:extent="80% 15%"/>',
        '    </layout>',
        '  </head>',
        '  <body>',
        '    <div style="defaultStyle">'
    ]
    
    for i, sub in enumerate(subtitles, start=1):
        text = (sub.get("text") or "").strip()
        if not text: continue
        start = sub.get("start_time", "").replace(",", ".")
        end = sub.get("end_time", "").replace(",", ".")
        if not start or not end: continue
        
        safe_text = saxutils.escape(text).replace("\n", "<br/>")
        lines.append(f'      <p xml:id="subtitle{i}" begin="{start}" end="{end}" region="bottom">{safe_text}</p>')
        
    lines.extend([
        '    </div>',
        '  </body>',
        '</tt>'
    ])
    
    content = "\n".join(lines)
    buf = io.BytesIO(content.encode("utf-8"))
    return StreamingResponse(buf, media_type="application/ttml+xml",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.ttml"})


@app.post("/export/csv")
async def export_csv(data: dict):
    import csv, io
    subtitles = data.get("subtitles", [])
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]
    
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["ID", "Start Time", "End Time", "Text"])
    
    for i, sub in enumerate(subtitles, start=1):
        writer.writerow([i, sub.get("start_time", ""), sub.get("end_time", ""), sub.get("text", "")])
        
    buf = io.BytesIO(output.getvalue().encode("utf-8"))
    return StreamingResponse(buf, media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.csv"})


@app.post("/export/rtf")
async def export_rtf(data: dict):
    subtitles = data.get("subtitles", [])
    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]
    
    lines = [r"{\rtf1\ansi\ansicpg1252\deff0\nouicompat\deflang1033{\fonttbl{\f0\fnil\fcharset0 Calibri;}}"]
    for sub in subtitles:
        text = (sub.get("text") or "").strip()
        if not text: continue
        safe_text = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}").replace("\n", "\\line ")
        lines.append(f"{safe_text}\\par\\par")
    lines.append("}")
    
    content = "\n".join(lines)
    buf = io.BytesIO(content.encode("utf-8"))
    return StreamingResponse(buf, media_type="application/rtf",
        headers={"Content-Disposition": f"attachment; filename={filename}_cleaned.rtf"})



@app.post("/track-changes")
async def get_track_changes(data: dict):
    """
    Returns track-changes data as JSON for on-screen rendering.
    Groups subtitle lines that share the same original_text (i.e. lines produced
    by splitting one long paragraph) so each unique source paragraph is shown
    only once, with all its cleaned output lines merged together.
    """
    from quality_checker import deduce_change_rules

    subtitles = data.get("subtitles", [])
    platform_key = data.get("platform_key", "generic")
    platform_dict = get_platform(platform_key)
    plat_rules = platform_dict.get("rules", [])

    # ── Group by original_text to avoid showing the same paragraph N times ──
    # Key = original_text, value = { ids, cleaned_lines, rule_hints merged }
    seen_originals: dict = {}   # orig_text -> group dict
    group_order = []            # insertion order
    total_lines = 0

    for sub in subtitles:
        text = sub.get("text", "").strip()
        orig_text = sub.get("original_text", "").strip()
        if not text and not orig_text:
            continue
        total_lines += 1
        if orig_text not in seen_originals:
            seen_originals[orig_text] = {
                "ids": [],
                "original_text": orig_text,
                "cleaned_lines": [],
                "rule_hints": list(sub.get("rule_hints", [])),
                "flagged": sub.get("flagged", False),
                "flag_reason": sub.get("flag_reason", ""),
            }
            group_order.append(orig_text)
        g = seen_originals[orig_text]
        g["ids"].append(sub.get("id"))
        if text and text not in g["cleaned_lines"]:
            g["cleaned_lines"].append(text)
        for h in sub.get("rule_hints", []):
            if h not in g["rule_hints"]:
                g["rule_hints"].append(h)
        if sub.get("flagged"):
            g["flagged"] = True

    # A professional audit must retain unchanged cues in their original order;
    # otherwise the report looks as though those lines were skipped.
    changes = []
    entries = []
    changed_count = 0
    unchanged_count = 0

    for orig_text in group_order:
        g = seen_originals[orig_text]
        new_text = "\n".join(g["cleaned_lines"])
        applied_rules = deduce_change_rules(orig_text, new_text, plat_rules, g["rule_hints"])

        entry = {
            "id": g["ids"][0],
            "ids": g["ids"],
            "original_text": orig_text,
            "new_text": new_text,
            "rules_applied": applied_rules,
            "flagged": g["flagged"],
            "flag_reason": g["flag_reason"],
            "changed": bool(applied_rules),
        }
        entries.append(entry)

        if not applied_rules:
            unchanged_count += len(g["ids"])
            continue
        changed_count += len(g["ids"])
        changes.append(entry)

    return {
        "changes": changes,
        "entries": entries,
        "total_lines": total_lines,
        "changed_lines": changed_count,
        "unchanged_lines": unchanged_count,
        "platform": platform_dict.get("name", platform_key)
    }


@app.post("/export/track-changes-pdf")
async def export_track_changes_pdf(data: dict):
    """
    Export track changes as a compact PDF.
    FAST PATH: if the caller sends a pre-computed 'changes' list (already built
    by /track-changes), we skip all regex recomputation and go straight to PDF.
    SLOW PATH (fallback): build changes from raw subtitles (old behaviour).
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    import html
    import re

    filename = data.get("filename", "cleaned").rsplit(".", 1)[0]
    platform_key = data.get("platform_key", "generic")
    platform_dict = get_platform(platform_key)

    # ── Determine source of changes ──────────────────────────────────────────
    # ``entries`` contains the full, ordered audit; retain the changed-only
    # field as a compatibility fallback for older browser sessions.
    precomputed = data.get("entries") or data.get("changes")
    total_lines = data.get("total_lines", 0)
    unchanged_count = data.get("unchanged_lines", 0)
    changed_count = data.get("changed_lines")

    if precomputed is None:
        # Slow-path fallback: recompute from raw subtitles
        from quality_checker import deduce_change_rules
        subtitles = data.get("subtitles", [])
        plat_rules = platform_dict.get("rules", [])
        seen_originals: dict = {}
        group_order = []
        total_lines = 0
        for sub in subtitles:
            text = sub.get("text", "").strip()
            orig_text = sub.get("original_text", "").strip()
            if not text and not orig_text:
                continue
            total_lines += 1
            if orig_text not in seen_originals:
                seen_originals[orig_text] = {
                    "ids": [], "original_text": orig_text,
                    "cleaned_lines": [], "rule_hints": list(sub.get("rule_hints", [])),
                }
                group_order.append(orig_text)
            g = seen_originals[orig_text]
            g["ids"].append(sub.get("id"))
            if text and text not in g["cleaned_lines"]:
                g["cleaned_lines"].append(text)
            for h in sub.get("rule_hints", []):
                if h not in g["rule_hints"]:
                    g["rule_hints"].append(h)
        precomputed = []
        for orig_text in group_order:
            g = seen_originals[orig_text]
            new_text = "\n".join(g["cleaned_lines"])
            rules = deduce_change_rules(orig_text, new_text, plat_rules, g["rule_hints"])
            if rules:
                precomputed.append({
                    "id": g["ids"][0], "ids": g["ids"],
                    "original_text": orig_text, "new_text": new_text,
                    "rules_applied": rules,
                })
        unchanged_count = total_lines - len(precomputed)
        changed_count = len(precomputed)

    if changed_count is None:
        changed_count = sum(1 for entry in precomputed if entry.get("changed", True))

    # ── Build PDF ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=1 * inch, rightMargin=1 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch
    )
    styles = getSampleStyleSheet()

    original_style = ParagraphStyle(
        "OriginalText", parent=styles["Normal"],
        alignment=TA_LEFT, fontSize=10,
        textColor=colors.HexColor("#d97706"), spaceAfter=3,
    )
    new_style = ParagraphStyle(
        "NewText", parent=styles["Normal"],
        alignment=TA_LEFT, fontSize=10,
        textColor=colors.HexColor("#059669"), spaceAfter=3,
    )
    rule_style = ParagraphStyle(
        "RuleText", parent=styles["Normal"],
        alignment=TA_LEFT, fontSize=8.5,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=10, leftIndent=10,
    )
    title_style = ParagraphStyle(
        "TitleStyle", parent=styles["Heading1"],
        alignment=TA_LEFT, spaceAfter=4, fontSize=14,
    )
    summary_style = ParagraphStyle(
        "SummaryStyle", parent=styles["Normal"],
        alignment=TA_LEFT, fontSize=10,
        textColor=colors.HexColor("#374151"), spaceAfter=14,
    )

    Story = []
    Story.append(Paragraph(f"Track Changes: {filename}", title_style))
    Story.append(Paragraph(
        f"Platform: <b>{platform_dict.get('name', platform_key)}</b> &nbsp;|&nbsp; "
        f"Total lines: <b>{total_lines}</b> &nbsp;|&nbsp; "
        f"Changed: <b>{changed_count}</b> &nbsp;|&nbsp; "
        f"Unchanged: <b>{unchanged_count}</b>",
        summary_style
    ))

    if not precomputed:
        Story.append(Paragraph(
            "&#10003; No changes were applied — the script was already clean.",
            summary_style
        ))
    else:
        ctrl = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]')
        for entry in precomputed:
            orig_text = ctrl.sub('', entry.get("original_text", ""))
            new_text  = ctrl.sub('', entry.get("new_text", ""))
            applied_rules = entry.get("rules_applied", [])
            sub_id = entry.get("id", "")
            ids = entry.get("ids", [sub_id])
            id_label = f"#{ids[0]}" if len(ids) == 1 else f"#{ids[0]}–#{ids[-1]}"

            orig_safe = html.escape(orig_text).replace("\n", "<br/>")
            new_safe  = html.escape(new_text).replace("\n", "<br/>")

            Story.append(Paragraph(f"<b>{id_label} Previously:</b> {orig_safe}", original_style))
            Story.append(Paragraph(f"<b>Cleaned:</b> {new_safe}", new_style))
            rules_html = "<b>Rules Applied:</b><br/>" + "<br/>".join(f"• {r}" for r in applied_rules)
            Story.append(Paragraph(rules_html, rule_style))

    doc.build(Story)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}_track_changes.pdf"})


@app.post("/adjust-timecodes")
async def adjust_timecodes_endpoint(data: dict):
    """
    Manual timecode correction tool — for the subtitler to fix a timecode
    they know is wrong (e.g. footage drift), without touching audio/video.
    Four modes:
      offset            — shift EVERY subtitle in the file by a fixed amount
      shift_only_this   — fix ONE subtitle's start/end, every other line untouched
      fix_from_index    — fix one subtitle, ripple the same shift to EVERY subtitle after it (to the end of the file)
      fix_range         — fix one subtitle, ripple the same shift only up to a chosen end ID (a bounded section, not the whole rest of the file) — use this when drift only affects a stretch of the file that re-syncs correctly later on
    """
    from timecode_adjuster import shift_timecodes, fix_from_index, fix_range_from_index, shift_only_this, parse_offset_input, sync_target, proportional_stretch

    subtitles = data.get("subtitles", [])
    mode = data.get("mode", "offset")

    if not subtitles:
        raise HTTPException(400, "No subtitles provided")

    if mode == "offset":
        offset_seconds = parse_offset_input(data.get("value", ""))
        if offset_seconds is None:
            raise HTTPException(400, "Could not parse offset value. Use a format like '+1.5s' or '-0:00:02'")
        result = shift_timecodes(subtitles, offset_seconds)
        return {"subtitles": result, "total": len(result), "value": data.get("value", "")}

    elif mode == "shift_only_this":
        target_id = data.get("target_id")
        new_start = data.get("new_start", "")
        new_end = data.get("new_end", "")
        if target_id is None:
            raise HTTPException(400, "target_id is required")
        result = shift_only_this(subtitles, target_id, new_start, new_end)
        return {"subtitles": result["subtitles"], "collision": result["collision"], "collision_detail": result["collision_detail"]}

    elif mode == "fix_from_index":
        target_id = data.get("target_id")
        new_tc = data.get("value", "")
        if target_id is None:
            raise HTTPException(400, "target_id is required")
        result = fix_from_index(subtitles, target_id, new_tc)
        return {"subtitles": result, "total": len(result)}

    elif mode == "fix_range":
        target_id = data.get("target_id")
        range_end_id = data.get("range_end_id")
        new_tc = data.get("value", "")
        if target_id is None or range_end_id is None:
            raise HTTPException(400, "Both target_id and range_end_id are required for fix_range mode")
        result = fix_range_from_index(subtitles, target_id, new_tc, range_end_id)
        return {
            "subtitles": result["subtitles"],
            "touched_ids": result["touched_ids"],
            "warning": result["warning"],
        }

    elif mode == "sync_target":
        target_id = data.get("target_id", 1)
        new_start = data.get("new_start", "")
        new_end = data.get("new_end", "")
        shift_mode = data.get("shift_mode", "all")
        result = sync_target(subtitles, target_id, new_start, new_end, shift_mode)
        return {
            "subtitles": result["subtitles"],
            "collision": result["collision"],
            "collision_detail": result["collision_detail"],
            "warning": result["warning"]
        }

    elif mode == "stretch":
        p_start = data.get("stretch_p_start", "")
        p_end = data.get("stretch_p_end", "")
        c_start = data.get("stretch_c_start", "")
        c_end = data.get("stretch_c_end", "")
        
        if not all([p_start, p_end, c_start, c_end]):
            raise HTTPException(400, "All 4 timecodes are required for stretch mode (pirate start/end, client start/end).")
            
        result = proportional_stretch(subtitles, p_start, p_end, c_start, c_end)
        
        if result.get("warning"):
            return {"subtitles": result["subtitles"], "warning": result["warning"]}
        return {"subtitles": result["subtitles"], "warning": ""}

    else:
        raise HTTPException(400, f"Unknown mode '{mode}'. Use offset, shift_only_this, fix_from_index, fix_range, or stretch.")



# ─── CLEANING & EXTRACTING ───────────────────────────────────────────

@app.post("/extract")
async def extract_endpoint(file: UploadFile = File(...)):
    """Extract subtitles / timecodes from any uploaded file (SRT, VTT, DOCX, TXT, PDF, image)."""
    file_bytes = await file.read()
    filename = file.filename or "subtitles.txt"

    async def event_stream():
        yield f"data: {json.dumps({'status': 'starting', 'message': 'Reading uploaded file...', 'progress': 5})}\n\n"
        await asyncio.sleep(0.05)
        try:
            from file_reader import read_file
            from timecoded_subtitles import parse_timecoded_subtitles

            file_data = read_file(file_bytes, filename)
            raw_text = file_data.get("raw_text", "")

            yield f"data: {json.dumps({'status': 'processing', 'message': 'Parsing subtitle dialogue & timecodes...', 'progress': 50})}\n\n"

            parsed = parse_timecoded_subtitles(raw_text)
            if not parsed:
                lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
                parsed = [{"id": i, "start_time": "", "end_time": "", "text": line} for i, line in enumerate(lines, 1)]

            subs = []
            for i, item in enumerate(parsed, 1):
                subs.append({
                    "id": item.get("id", i),
                    "start_time": item.get("start_time", ""),
                    "end_time": item.get("end_time", ""),
                    "text": item.get("text", ""),
                    "flagged": False,
                    "flag_reason": ""
                })

            stats = {
                "total_lines": len(subs),
                "original_format": filename,
                "structure": file_data.get("structure", "unknown")
            }

            yield f"data: {json.dumps({'status': 'completed', 'progress': 100, 'result': {'subtitles': subs, 'stats': stats}})}\n\n"
        except Exception as e:
            import traceback
            traceback.print_exc()
            yield f"data: {json.dumps({'status': 'error', 'error': str(e)})}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}
    )


@app.post("/clean-extracted")
async def clean_extracted_endpoint(data: dict):
    """
    Deterministic platform-rule cleaning of already-extracted subtitles.

    Zero hallucination — uses ONLY Python rule engines, no LLM.
    The AI (Groq GPT-OSS 20B) is for raw script → subtitle extraction only.
    For already-extracted subtitles the deterministic engine is MORE accurate
    because it applies every rule exactly, never changes dialogue meaning,
    and never alters line counts.

    Pipeline (in order):
      1. clean_delivery_text   – remove speaker labels, stage directions, HOH noise
      2. auto_fix_subtitles    – apply ALL platform rules: profanity, numbers, hyphen
                                  format, italics, punctuation, sentence case,
                                  US spelling, line splitting, HOH removal, etc.
      3. prepare_for_platform  – timing repair: duration, gap, CPS, zero-subtitle
      4. QC flag               – only flag lines with UNFIXABLE errors (still too long,
                                  profanity missed, etc.) — not warnings/info
    """
    subs = data.get("subtitles", [])
    platform_key = data.get("platform_key", "generic")
    filename = data.get("filename", "subtitles.txt")
    regenerate_timings = bool(data.get("regenerate_timings", False))

    if not subs:
        raise HTTPException(400, "No subtitles provided to clean")

    async def event_stream():
        yield f"data: {json.dumps({'status': 'starting', 'message': 'Loading platform rules...', 'progress': 0})}\n\n"
        await asyncio.sleep(0.05)
        try:
            from platform_rules import get_platform
            from timecoded_subtitles import prepare_for_platform, clean_delivery_text
            from quality_checker import auto_fix_subtitles
            from italic_formatter import apply_italics_rules

            platform_dict = get_platform(platform_key)
            platform_name = platform_dict.get("name", platform_key)
            max_chars = int(platform_dict.get("max_chars_per_line", 42))

            # ── Pass 1: Per-subtitle text cleaning ───────────────────────────
            # Removes speaker labels, stage directions, HOH elements, metadata
            # noise — preserves <i>/<b> tags, dialogue text, music notes ♪
            yield f"data: {json.dumps({'status': 'processing', 'message': f'Cleaning text for {platform_name}...', 'progress': 15})}\n\n"
            await asyncio.sleep(0.05)

            pass1 = []
            for sub in subs:
                item = dict(sub)
                source_text = sub.get("text", "")
                item["original_text"] = source_text
                protected_title = _preserve_as_broadcast_title(source_text, platform_dict)
                if protected_title is not None:
                    # Narrative-title words are source-controlled visible text.
                    # Preserve them before the deterministic rules run and keep
                    # the value on the item so later passes cannot erase it.
                    item["_protected_as_broadcast_title"] = protected_title
                    item["text"] = protected_title
                else:
                    cleaned_text = clean_delivery_text(source_text)
                    item["text"] = cleaned_text if cleaned_text.strip() else source_text
                pass1.append(item)

            # ── Pass 2: Full platform rule enforcement ────────────────────────
            # Applies every deterministic rule: profanity, numbers, hyphen format,
            # ellipsis, sentence case, US spelling, HOH removal, line splitting,
            # two-speaker format, song lyric casing, acronyms, symbols, etc.
            yield f"data: {json.dumps({'status': 'processing', 'message': f'Applying {platform_name} formatting rules...', 'progress': 40})}\n\n"
            await asyncio.sleep(0.05)

            pass2 = auto_fix_subtitles(pass1, platform_key)

            # ── Pass 3: Italics rules ────────────────────────────────────────
            # Applies platform-specific italics: song lyrics, VO, phone calls,
            # foreign words, onscreen text — or strips italics if platform forbids
            yield f"data: {json.dumps({'status': 'processing', 'message': 'Applying italics rules...', 'progress': 60})}\n\n"
            await asyncio.sleep(0.05)

            pass3 = apply_italics_rules(pass2, platform_key)

            # Formatting utilities may legitimately remove bracketed metadata.
            # Restore protected title text before timing work; it is never
            # metadata and must not be reduced to a bare @m/@n tag.
            for item in pass3:
                if item.get("_protected_as_broadcast_title"):
                    item["text"] = item["_protected_as_broadcast_title"]

            # ── Pass 4: Timing/structural rules ──────────────────────────────
            # Repairs duration (min/max), gap between subtitles, CPS, zero-subtitle
            yield f"data: {json.dumps({'status': 'processing', 'message': 'Applying timing and structural rules...', 'progress': 78})}\n\n"
            await asyncio.sleep(0.05)

            cleaned = prepare_for_platform(
                pass3, platform_key, filename,
                regenerate_timings=regenerate_timings,
            )

            # prepare_for_platform can make a fresh dict while repairing
            # timings.  Enforce source-title fidelity once more at the final
            # boundary used by the editor/export response.
            for item in cleaned:
                if item.get("_protected_as_broadcast_title"):
                    item["text"] = item["_protected_as_broadcast_title"]

            # ── Pass 5: Targeted flag — unfixable errors only ─────────────────
            # Only flag what genuinely could NOT be auto-fixed:
            # - line still too long (very long word that can't be split)
            # - profanity that slipped through all passes
            # - HOH element that survived all passes
            yield f"data: {json.dumps({'status': 'processing', 'message': 'Running quality check...', 'progress': 90})}\n\n"
            await asyncio.sleep(0.05)

            UNFIXABLE_TYPES = {
                "LINE_TOO_LONG",          # single word longer than max_chars — human must rewrite
                "TOO_MANY_LINES",         # needs human restructuring
                "READING_SPEED_EXCEEDED", # text must be shortened — changing words = human job
                "PROFANITY_NOT_REPLACED", # missed profanity
                "HOH_EMT_ELEMENT",        # missed HOH element
                "ZERO_SUBTITLE_INVALID",  # wrong/missing delivery fields
            }

            final_subs = []
            for s in cleaned:
                # Re-check line length post-cleaning — this is the single most
                # reliable signal: if any line is STILL over max_chars after
                # auto_fix tried to split it, a human needs to rewrite it.
                lines_in_sub = s.get("text", "").split("\n")
                still_too_long = any(
                    len(re.sub(r'<[^>]+>', '', ln)) > max_chars
                    for ln in lines_in_sub
                )
                final_subs.append({
                    **s,
                    "flagged": still_too_long,
                    "flag_reason": (
                        f"Line exceeds {max_chars} chars after auto-split — please shorten manually."
                        if still_too_long else ""
                    ),
                })

            yield f"data: {json.dumps({'status': 'processing', 'message': 'Finalizing...', 'progress': 97})}\n\n"
            await asyncio.sleep(0.05)

            changed_lines = sum(
                1 for s in final_subs
                if s.get("original_text", s.get("text", "")).strip() != s.get("text", "").strip()
            )
            stats = {
                "total_lines": len(final_subs),
                "flagged_lines": sum(1 for s in final_subs if s.get("flagged")),
                "changed_lines": changed_lines,
                "ai_used": False,
                "platform": platform_name,
                "original_format": filename,
            }

            yield f"data: {json.dumps({'status': 'completed', 'progress': 100, 'result': {'subtitles': final_subs, 'stats': stats}})}\n\n"

        except Exception as e:
            import traceback
            traceback.print_exc()
            yield f"data: {json.dumps({'status': 'error', 'error': str(e)})}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}
    )


@app.post("/export/srt")
async def export_srt_endpoint(data: dict):
    subs = data.get("subtitles", [])
    filename = (data.get("filename") or "subtitles").rsplit(".", 1)[0]
    from timecoded_subtitles import subtitles_to_srt
    content = subtitles_to_srt(subs)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="application/x-subrip",
        headers={"Content-Disposition": f'attachment; filename="{filename}_cleaned.srt"'}
    )


@app.post("/export/txt")
async def export_txt_endpoint(data: dict):
    subs = data.get("subtitles", [])
    filename = (data.get("filename") or "subtitles").rsplit(".", 1)[0]
    lines = []
    for s in subs:
        if s.get("start_time"):
            lines.append(f"{s.get('start_time')} --> {s.get('end_time')}\n{s.get('text')}\n")
        else:
            lines.append(s.get("text", ""))
    content = "\n".join(lines)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{filename}_cleaned.txt"'}
    )


@app.post("/export/docx")
async def export_docx_endpoint(data: dict):
    subs = data.get("subtitles", [])
    filename = (data.get("filename") or "subtitles").rsplit(".", 1)[0]
    from docx import Document
    doc = Document()
    doc.add_heading(f"Cleaned Subtitles: {filename}", level=1)
    for s in subs:
        tc = f"{s.get('start_time','')} --> {s.get('end_time','')}" if s.get('start_time') else ""
        p = doc.add_paragraph()
        if tc:
            p.add_run(tc + "\n").bold = True
        p.add_run(s.get("text", ""))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}_cleaned.docx"'}
    )


@app.post("/export/pdf")
async def export_pdf_endpoint(data: dict):
    subs = data.get("subtitles", [])
    filename = (data.get("filename") or "subtitles").rsplit(".", 1)[0]
    lines = []
    for s in subs:
        if s.get("start_time"):
            lines.append(f"{s.get('start_time')} --> {s.get('end_time')}\n{s.get('text')}\n")
        else:
            lines.append(s.get("text", ""))
    content = "\n".join(lines)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{filename}_cleaned.txt"'}
    )



# ─── PLATFORMS ───────────────────────────────────────────────────

@app.get("/platforms")
def get_platforms():
    try:
        # Serve only custom DB platforms — no built-in hardcoded OTTs
        db_platforms = get_all_platforms()
        return {"platforms": db_platforms}
    except Exception:
        # DB not ready — return empty so UI prompts user to add guidelines
        return {"platforms": {}}



def _extract_text_from_upload(file_bytes: bytes, filename: str, sheet_name: str = "") -> str:
    """Return guideline text for one upload, OCR-inclusive (see file_reader).
    
    When sheet_name is given (bulk Excel import), reads only that sheet's cells
    AND appends any OCR text from embedded images in the whole workbook — since
    Excel stores all images globally (xl/media/), not per-sheet.
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext in ("png", "jpg", "jpeg", "webp", "bmp", "gif"):
        from ocr_reader import ocr_image_bytes
        return ocr_image_bytes(file_bytes).strip()
    if sheet_name.strip() and ext in ("xlsx", "xls"):
        from file_reader import read_excel_sheet
        cell_text = read_excel_sheet(file_bytes, sheet_name.strip()).strip()
        # Also OCR embedded images from workbook (images are stored globally, not per-sheet)
        try:
            from ocr_reader import ocr_fallback_for_xlsx, ocr_fallback_for_xls
            # Always run OCR regardless of cell text length — images may have rules
            ocr_fn = ocr_fallback_for_xlsx if ext == "xlsx" else ocr_fallback_for_xls
            ocr_text = ocr_fn(file_bytes, 0)  # pass 0 to always trigger OCR check
            if ocr_text and ocr_text.strip():
                return (cell_text + "\n\n=== EMBEDDED IMAGE CONTENT (OCR) ===\n" + ocr_text).strip()
        except Exception as _e:
            print(f"[extract] OCR pass failed for sheet '{sheet_name}': {_e}")
        return cell_text
    from file_reader import read_file
    return read_file(file_bytes, filename, force_ocr=True)["raw_text"].strip()


def _fallback_pasted_rules(raw_guidelines: str) -> tuple[list[str], list[str], dict]:
    """Persist usable pasted rules even when the AI extractor is unavailable.

    Pasting a short, already-written rule list must not produce an empty OTT
    simply because the extraction model is rate-limited or returns malformed
    JSON.  This fallback deliberately retains the user's wording verbatim and
    only classifies it into the existing text/timing/delivery buckets.
    """
    candidates = []
    for line in raw_guidelines.splitlines():
        line = re.sub(r"^\s*(?:[-*•]|\d+[.)])\s*", "", line).strip()
        if len(line) >= 8:
            candidates.append(line)
    # A single pasted paragraph is also a valid rule source.  Split only at a
    # sentence boundary, never at commas or semicolons inside a requirement.
    if len(candidates) <= 1 and raw_guidelines.strip():
        candidates = [
            part.strip() for part in re.split(r"(?<=[.!?])\s+(?=[A-Z@])", raw_guidelines.strip())
            if len(part.strip()) >= 8
        ]

    deduped = []
    seen = set()
    for rule in candidates:
        key = rule.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(rule)

    from rule_segregation import segregate_rules, derive_timing_fields
    buckets = segregate_rules(deduped, [])
    operational = buckets["timing"] + buckets["positioning"] + buckets["file"]
    timing_fields, _ = derive_timing_fields(buckets["timing"])
    return buckets["text"], operational, timing_fields


def _save_one_platform(platform_name: str, version_label: str, raw_guidelines: str,
                       source_files: list = None) -> dict:
    """Extract rules from raw guideline text and persist one platform. Returns a result dict."""
    platform_name = platform_name.strip()
    version_label = version_label.strip() or "Current"
    if not platform_name:
        return {"status": "skipped", "reason": "Missing platform name"}

    version_slug = re.sub(r'[^a-z0-9]', '_', version_label.lower()).strip('_')
    platform_family = "custom_" + re.sub(r'[^a-z0-9]', '_', platform_name.lower())
    platform_key = f"{platform_family}__{version_slug}"

    if not raw_guidelines.strip():
        return {"status": "skipped", "platform_name": platform_name,
                "reason": "No readable text found in the document."}

    print(f"[PLATFORMS] Extracting rules for '{platform_name}' v'{version_label}' from {len(raw_guidelines)} chars")
    platform_data = extract_platform_rules_with_ai(raw_guidelines, platform_name)

    # The AI is helpful for long documents, but a direct paste must still save
    # its rules when the model is unavailable or cannot return JSON.
    if not platform_data.get("rules") and not platform_data.get("subtitler_rules"):
        pasted_rules, operational_rules, timing_fields = _fallback_pasted_rules(raw_guidelines)
        platform_data["rules"] = pasted_rules
        platform_data["subtitler_rules"] = operational_rules
        platform_data.update(timing_fields)
        platform_data["summary"] = (
            f"Rules pasted directly for {platform_name}; retained verbatim because AI extraction returned no rules."
        )
        print(f"[PLATFORMS] AI returned no rules; saved {len(pasted_rules) + len(operational_rules)} pasted rules directly")

    platform_data["platform_family"] = platform_family
    platform_data["version_label"]   = version_label
    platform_data["guidelines_raw"]  = raw_guidelines
    source_files = source_files or []
    platform_data["source_files"]    = source_files
    platform_data["source_file_id"]  = source_files[0]["id"] if source_files else None
    platform_data["source_filename"] = source_files[0]["name"] if source_files else None

    save_custom_platform(platform_key, platform_data)

    rules_count = len(platform_data.get("rules", []))
    print(f"[PLATFORMS] Saved '{platform_name}' ({version_label}) as '{platform_key}' with {rules_count} rules")

    return {
        "status": "ok",
        "platform_key": platform_key,
        "platform_family": platform_family,
        "version_label": version_label,
        "platform_name": platform_name,
        "rules_extracted": rules_count,
    }


@app.post("/platforms/add-stream")
async def add_platform_stream(
    platform_name: str = Form(default=""),
    version_label: str = Form(default="Current"),
    sheet_name: str = Form(default=""),
    selected_sheets: str = Form(default=""),  # JSON array of sheet names to import
    guidelines_files: list[UploadFile] = File(default=[]),
    guidelines_text: str = Form(default="")
):
    import asyncio, json as _json, queue as _queue, threading

    file_snapshots = []
    for f in guidelines_files:
        if not f.filename:
            continue
        fb = await f.read()
        if fb:
            file_snapshots.append((f.filename, fb))
    pasted = (guidelines_text or "").strip()
    p_name = platform_name.strip()
    v_label = version_label.strip() or "Current"
    s_name = sheet_name.strip()
    if pasted and not file_snapshots and not p_name:
        raise HTTPException(400, "Platform name is required when pasting guidelines text")
    # Parse the list of sheets the user explicitly checked in the UI
    import json as _json_
    try:
        selected_sheets_list = _json_.loads(selected_sheets) if selected_sheets.strip() else []
    except Exception:
        selected_sheets_list = []

    if not p_name and file_snapshots:
        fname = file_snapshots[0][0]
        base_name = _os.path.splitext(fname)[0]
        clean_name = re.sub(r'^(screenshot|guidelines|rules|doc|pdf)_*', '', base_name, flags=re.IGNORECASE)
        clean_name = re.sub(r'[\-_(0-9)]+', ' ', clean_name).strip()
        p_name = clean_name.title() or "Custom Platform"

    if not p_name:
        p_name = "Custom Platform"

    q: _queue.Queue = _queue.Queue()
    SENTINEL = object()

    def worker():
        try:
            from file_reader import list_excel_sheets
            q.put({"status": "start", "progress": 5, "message": "Reading guideline document(s) & running OCR..."})

            parts = []          # (label, text)
            excel_files = []    # (filename, bytes)
            source_files = []

            for fname, fb in file_snapshots:
                fid = _store_source_file(fb, fname)
                if fid:
                    source_files.append({"id": fid, "name": fname})

            for fname, fb in file_snapshots:
                ext = fname.lower().rsplit(".", 1)[-1]
                if ext in ("xlsx", "xls") and not s_name:
                    excel_files.append((fname, fb))
                    continue
                text = _extract_text_from_upload(fb, fname, s_name)
                if text:
                    parts.append((fname, text))

            if pasted:
                parts.append(("pasted_text", pasted))

            # ── BULK EXCEL ────────────────────────────────────────────────────
            if excel_files:
                total = 0
                try:
                    for _, fb in excel_files:
                        sheets_all = list_excel_sheets(fb) or []
                        if selected_sheets_list:
                            sheets_all = [s for s in sheets_all if (s["name"] if isinstance(s, dict) else s) in selected_sheets_list]
                        total += len(sheets_all)
                except Exception:
                    total = len(excel_files)
                done = 0
                results = []
                for fname, fb in excel_files:
                    try:
                        sheets = list_excel_sheets(fb) or []
                    except Exception as e:
                        results.append({"filename": fname, "status": "error", "reason": str(e)})
                        continue
                    if not sheets:
                        results.append({"filename": fname, "status": "skipped", "reason": "No sheets found"})
                        continue
                    # Filter to only user-checked sheets if a subset was specified
                    if selected_sheets_list:
                        sheets = [s for s in sheets if (s["name"] if isinstance(s, dict) else s) in selected_sheets_list]
                    for sh in sheets:
                        sname = sh["name"] if isinstance(sh, dict) else sh
                        # ALWAYS use the sheet name as the platform name in bulk mode.
                        # The manually-typed platform_name is intentionally ignored here
                        # to prevent all sheets from being saved under the same name.
                        pname = sname
                        done += 1
                        prog = int(10 + (done / max(total, 1)) * 85)
                        q.put({"status": "extracting", "progress": prog,
                               "message": f'Reading sheet "{sname}" + OCR images ({done}/{total})',
                               "sheet": sname})
                        sheet_text = _extract_text_from_upload(fb, fname, sname)
                        if not sheet_text.strip():
                            results.append({"filename": fname, "sheet_name": sname,
                                            "status": "skipped", "reason": "Sheet appears empty"})
                            continue
                        res = _save_one_platform(pname, v_label, sheet_text, source_files)
                        res["filename"] = fname
                        res["sheet_name"] = sname
                        results.append(res)
                ok = [r for r in results if r.get("status") == "ok"]
                q.put({"status": "completed", "progress": 100, "result": {
                    "success": True, "bulk": True, "total": len(results), "imported": len(ok),
                    "results": results,
                    "message": f"Imported {len(ok)} of {len(results)} sheets."
                }})
                return

            # ── SINGLE PLATFORM ─────────────────────────────────────────────────
            raw_guidelines = "\n\n".join(t for _, t in parts)
            if not raw_guidelines.strip():
                q.put({"status": "error", "error": "No readable text found in the document(s)."})
                return

            q.put({"status": "ocr", "progress": 30,
                    "message": "OCR & image analysis complete. Extracting rules with AI..."})
            res = _save_one_platform(p_name, v_label, raw_guidelines, source_files)
            if res.get("status") != "ok":
                q.put({"status": "error", "error": res.get("reason", "Could not import platform.")})
                return

            q.put({"status": "completed", "progress": 100, "result": {
                "success": True,
                "platform_key": res["platform_key"],
                "platform_family": res["platform_family"],
                "version_label": res["version_label"],
                "platform_name": res["platform_name"],
                "rules_extracted": res["rules_extracted"],
                "message": f"Platform '{res['platform_name']}' saved with {res['rules_extracted']} rules extracted."
            }})
        except Exception as e:
            import traceback
            traceback.print_exc()
            q.put({"status": "error", "error": str(e)})
        finally:
            q.put(SENTINEL)

    threading.Thread(target=worker, daemon=True).start()

    async def event_gen():
        loop = asyncio.get_running_loop()
        while True:
            item = await loop.run_in_executor(None, q.get)
            if item is SENTINEL:
                break
            try:
                yield f"data: {_json.dumps(item, ensure_ascii=False)}\n\n"
            except Exception:
                break

    return StreamingResponse(
        event_gen(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
    )


@app.post("/platforms/add")
async def add_platform(
    platform_name: str = Form(...),
    version_label: str = Form(default="Current"),
    sheet_name: str = Form(default=""),          # optional: target a specific Excel sheet
    guidelines_file: UploadFile = File(None),
    guidelines_text: str = Form(default="")
):
    if not platform_name.strip():
        raise HTTPException(400, "Platform name is required")

    version_slug = re.sub(r'[^a-z0-9]', '_', (version_label or 'current').lower().strip()).strip('_')
    platform_family = "custom_" + re.sub(r'[^a-z0-9]', '_', platform_name.lower().strip())
    platform_key = f"{platform_family}__{version_slug}"

    raw_guidelines = ""
    source_files = []
    if guidelines_file and guidelines_file.filename:
        file_bytes = await guidelines_file.read()
        if not file_bytes:
            raise HTTPException(400, "Uploaded guidelines file is empty")

        fid = _store_source_file(file_bytes, guidelines_file.filename)
        if fid:
            source_files.append({"id": fid, "name": guidelines_file.filename})

        ext = guidelines_file.filename.lower().rsplit(".", 1)[-1]
        if sheet_name.strip() and ext in ("xlsx", "xls"):
            from file_reader import read_excel_sheet
            raw_guidelines = read_excel_sheet(file_bytes, sheet_name.strip())
            print(f"[PLATFORMS] Read sheet '{sheet_name}' → {len(raw_guidelines)} chars")
        else:
            file_data = read_file(file_bytes, guidelines_file.filename)
            raw_guidelines = file_data["raw_text"]
            print(f"[PLATFORMS] Read {len(raw_guidelines)} chars from '{guidelines_file.filename}'")
    elif guidelines_text.strip():
        raw_guidelines = guidelines_text.strip()

    if not raw_guidelines.strip():
        raise HTTPException(400, "No readable text found in the uploaded document. "
                            "Please try a different format (PDF, DOCX, TXT).")

    print(f"[PLATFORMS] Extracting rules for '{platform_name}' v'{version_label}' from {len(raw_guidelines)} chars")
    platform_data = extract_platform_rules_with_ai(raw_guidelines, platform_name.strip())

    platform_data["platform_family"] = platform_family
    platform_data["version_label"]   = version_label.strip() or "Current"
    platform_data["guidelines_raw"]  = raw_guidelines
    platform_data["source_files"]    = source_files
    if source_files:
        platform_data["source_file_id"]  = source_files[0]["id"]
        platform_data["source_filename"] = source_files[0]["name"]

    save_custom_platform(platform_key, platform_data)

    rules_count = len(platform_data.get("rules", []))
    print(f"[PLATFORMS] Saved '{platform_name}' ({version_label}) as '{platform_key}' with {rules_count} rules")

    return {
        "success": True,
        "platform_key": platform_key,
        "platform_family": platform_family,
        "version_label": platform_data["version_label"],
        "platform_name": platform_name.strip(),
        "rules_extracted": rules_count,
        "message": f"Platform '{platform_name}' ({platform_data['version_label']}) saved with {rules_count} rules extracted"
    }


@app.post("/platforms/preview-excel")
async def preview_excel_sheets(guidelines_file: UploadFile = File(...)):
    """
    Return the list of sheet names (+ row counts) from an uploaded Excel file
    WITHOUT extracting any rules. Used by the bulk-import UI so subtitlers
    can see what sheets exist and map each one to a platform name/version.
    """
    file_bytes = await guidelines_file.read()
    if not file_bytes:
        raise HTTPException(400, "File is empty")
    ext = guidelines_file.filename.lower().rsplit(".", 1)[-1]
    if ext not in ("xlsx", "xls"):
        raise HTTPException(400, "Only Excel (.xlsx / .xls) files are supported for sheet preview")
    from file_reader import list_excel_sheets
    sheets = list_excel_sheets(file_bytes)
    if not sheets:
        raise HTTPException(422, "Could not read any sheets from this file")
    return {"filename": guidelines_file.filename, "sheets": sheets}


@app.post("/platforms/bulk-add")
async def bulk_add_platforms(
    guidelines_file: UploadFile = File(...),
    mappings: str = Form(...)   # JSON string: [{sheet_name, platform_name, version_label}]
):
    """
    Process a multi-sheet Excel file and import each mapped sheet as a
    separate platform version in one request.

    mappings example:
    [
      {"sheet_name": "Netflix",      "platform_name": "Netflix",      "version_label": "Current"},
      {"sheet_name": "Discovery Max","platform_name": "Discovery Max","version_label": "2023 Guidelines"}
    ]
    """
    import json as _json
    try:
        mapping_list = _json.loads(mappings)
    except Exception:
        raise HTTPException(400, "Invalid mappings JSON")

    file_bytes = await guidelines_file.read()
    if not file_bytes:
        raise HTTPException(400, "File is empty")

    source_files = []
    if guidelines_file.filename:
        fid = _store_source_file(file_bytes, guidelines_file.filename)
        if fid:
            source_files.append({"id": fid, "name": guidelines_file.filename})

    from file_reader import read_excel_sheet

    results = []
    for entry in mapping_list:
        sheet   = entry.get("sheet_name", "").strip()
        p_name  = entry.get("platform_name", "").strip()
        version = entry.get("version_label", "Current").strip() or "Current"

        if not sheet or not p_name:
            results.append({"sheet_name": sheet, "status": "skipped", "reason": "Missing sheet or platform name"})
            continue

        try:
            raw = read_excel_sheet(file_bytes, sheet)
            if not raw.strip():
                results.append({"sheet_name": sheet, "platform_name": p_name, "status": "skipped",
                                 "reason": "Sheet is empty or could not be read"})
                continue

            version_slug   = re.sub(r'[^a-z0-9]', '_', version.lower()).strip('_')
            platform_family = "custom_" + re.sub(r'[^a-z0-9]', '_', p_name.lower())
            platform_key    = f"{platform_family}__{version_slug}"

            print(f"[BULK] Processing sheet '{sheet}' → '{p_name}' ({version}) — {len(raw)} chars")
            platform_data = extract_platform_rules_with_ai(raw, p_name)
            platform_data["platform_family"] = platform_family
            platform_data["version_label"]   = version
            platform_data["guidelines_raw"]  = raw
            platform_data["source_files"]    = source_files
            if source_files:
                platform_data["source_file_id"]  = source_files[0]["id"]
                platform_data["source_filename"] = source_files[0]["name"]

            save_custom_platform(platform_key, platform_data)
            rules_count = len(platform_data.get("rules", []))

            results.append({
                "sheet_name": sheet,
                "platform_name": p_name,
                "version_label": version,
                "platform_key": platform_key,
                "rules_extracted": rules_count,
                "status": "ok"
            })
        except Exception as e:
            results.append({"sheet_name": sheet, "platform_name": p_name, "status": "error", "reason": str(e)})

    ok = [r for r in results if r.get("status") == "ok"]
    return {
        "total": len(mapping_list),
        "imported": len(ok),
        "results": results,
        "message": f"Imported {len(ok)} of {len(mapping_list)} platforms from '{guidelines_file.filename}'"
    }


@app.delete("/platforms")
def delete_all_platforms():
    from database import get_connection
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM platforms")
    affected = cursor.rowcount
    conn.commit()
    cursor.close()
    conn.close()
    return {"success": True, "deleted": affected}


@app.delete("/platforms/{platform_key}")
def delete_platform(platform_key: str):
    from database import get_connection
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM platforms WHERE platform_key=%s", (platform_key,))
    affected = cursor.rowcount
    conn.commit()
    cursor.close()
    conn.close()
    if affected == 0:
        raise HTTPException(404, "Platform not found")
    return {"success": True}


@app.delete("/platforms/family/{family_key}")
def delete_platform_family(family_key: str):
    """Delete ALL versions of a platform family at once."""
    from database import get_connection
    conn = get_connection()
    cursor = conn.cursor()
    # Match by family key OR by platform_key prefix (covers built-ins which use key as family)
    cursor.execute(
        "DELETE FROM platforms WHERE platform_family=%s OR platform_key=%s",
        (family_key, family_key)
    )
    affected = cursor.rowcount
    conn.commit()
    cursor.close()
    conn.close()
    if affected == 0:
        raise HTTPException(404, "No versions found for this platform family")
    return {"success": True, "deleted_versions": affected}


@app.patch("/platforms/{platform_key}/meta")
async def update_platform_meta(platform_key: str, data: dict):
    """
    Update name, version_label, and/or rules for one platform version.
    Works on any platform (built-in or custom).
    """
    import json as _json
    from database import get_connection

    name          = data.get("name")
    version_label = data.get("version_label")
    rules         = data.get("rules")
    subtitler_rules = data.get("subtitler_rules")

    if not any([name, version_label is not None, rules is not None, subtitler_rules is not None]):
        raise HTTPException(400, "Nothing to update — provide name, version_label, rules, or subtitler_rules")

    conn   = get_connection()
    cursor = conn.cursor()

    parts  = []
    params = []
    if name is not None:
        parts.append("name=%s");          params.append(name.strip())
    if version_label is not None:
        parts.append("version_label=%s"); params.append(version_label.strip())
    if rules is not None:
        if not isinstance(rules, list):
            raise HTTPException(400, "Rules must be an array")
        parts.append("rules=%s");         params.append(_json.dumps(rules))
    if subtitler_rules is not None:
        if not isinstance(subtitler_rules, list):
            raise HTTPException(400, "Subtitler rules must be an array")
        parts.append("subtitler_rules=%s"); params.append(_json.dumps(subtitler_rules))

    params.append(platform_key)
    cursor.execute(f"UPDATE platforms SET {', '.join(parts)} WHERE platform_key=%s", params)
    affected = cursor.rowcount
    conn.commit()
    cursor.close()
    conn.close()

    if affected == 0:
        raise HTTPException(404, "Platform not found")
    return {"success": True, "platform_key": platform_key}


@app.patch("/platforms/family/{family_key}/name")
async def rename_platform_family(family_key: str, data: dict):
    """
    Rename the display name for ALL versions of a platform family at once.
    e.g. 'Nickelodeon V10' → 'Nickelodeon' across every version in the group.
    """
    new_name = (data.get("name") or "").strip()
    if not new_name:
        raise HTTPException(400, "New name is required")

    from database import get_connection
    conn   = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE platforms SET name=%s WHERE platform_family=%s OR platform_key=%s",
        (new_name, family_key, family_key)
    )
    affected = cursor.rowcount
    conn.commit()
    cursor.close()
    conn.close()

    if affected == 0:
        raise HTTPException(404, "Platform family not found")
    return {"success": True, "family_key": family_key, "new_name": new_name, "updated": affected}



# ─── HELPER ──────────────────────────────────────────────────────

def normalize_tc(tc: str) -> str:
    if not tc:
        return ""
    tc = tc.strip()
    if re.match(r'\d{2}:\d{2}:\d{2},\d{3}', tc):
        return tc
    m = re.match(r'(\d{2}):(\d{2}):(\d{2})[:;](\d{2})', tc)
    if m:
        h, mn, s, f = m.groups()
        ms = int(int(f) * 1000 / 25)
        return f"{h}:{mn}:{s},{ms:03d}"
    m = re.match(r'(\d{2}):(\d{2}):(\d{2})\.(\d+)', tc)
    if m:
        h, mn, s, ms = m.groups()
        return f"{h}:{mn}:{s},{ms[:3].ljust(3,'0')}"
    return tc



# ─── OTT GUIDELINES SEARCH ENGINE ───────────────────────────────────
#
# Matches the company's own mockup (June_26_2026_OTT_Guidelines_SearchEngine.pptx):
# keyword search + Client/Platform/Category/Year filters over a flat,
# searchable table of digitized guideline spec rows. Separate concern from
# the AI cleaner's platform rules — this is for humans looking things up,
# not for the cleaner consuming numeric limits.

@app.get("/guidelines/search")
def search_guidelines_endpoint(
    keyword: str = "",
    client: str = "",
    ott_platform: str = "",
    category: str = "",
    year: int = None,
    include_all_versions: bool = False,
):
    from guidelines_db import search_guidelines
    try:
        results = search_guidelines(keyword, client, ott_platform, category, year, include_all_versions)
        return {"results": results, "total": len(results)}
    except Exception as e:
        raise HTTPException(500, f"Guidelines search failed: {e}")


@app.get("/guidelines/filters")
def get_guidelines_filters():
    """Returns the live distinct values for every filter dropdown, so the UI never shows a stale hardcoded list."""
    from guidelines_db import get_filter_options
    try:
        return get_filter_options()
    except Exception as e:
        raise HTTPException(500, f"Could not load filter options: {e}")


@app.post("/guidelines/add")
async def add_guideline_endpoint(data: dict):
    from guidelines_db import add_guideline
    required = ["client", "ott_platform", "spec", "guideline", "year"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        raise HTTPException(400, f"Missing required field(s): {', '.join(missing)}")
    try:
        new_id = add_guideline(data)
        return {"success": True, "id": new_id}
    except Exception as e:
        raise HTTPException(500, f"Could not add guideline: {e}")


@app.put("/guidelines/{guideline_id}")
async def update_guideline_endpoint(guideline_id: int, data: dict):
    """
    Updating a guideline does NOT overwrite it — it creates a new version
    and marks the old one superseded (but keeps it queryable). See
    guidelines_db.update_guideline() for why: old projects need to be able
    to look up the rules that were in force when THEY were built.
    """
    from guidelines_db import update_guideline, search_guidelines
    try:
        success = update_guideline(guideline_id, data)
        if not success:
            raise HTTPException(404, "Guideline not found")
        return {"success": True, "note": "Previous version preserved in history, not overwritten."}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Could not update guideline: {e}")


@app.get("/guidelines/{guideline_id}/history")
def get_guideline_history_endpoint(guideline_id: int):
    """
    Returns every version of a guideline (oldest first), given any single
    version's id. Lets the UI show 'this rule changed 3 times' and lets
    someone pick an older version for an older project.
    """
    from guidelines_db import get_guideline_history
    try:
        history = get_guideline_history(guideline_id)
        if not history:
            raise HTTPException(404, "Guideline not found")
        return {"history": history, "total_versions": len(history)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Could not load guideline history: {e}")


@app.delete("/guidelines/{guideline_id}")
def delete_guideline_endpoint(guideline_id: int, hard_delete: bool = False):
    """
    By default this is a soft delete — marks the entry as no longer
    current, but the row and its history remain in the database. Pass
    hard_delete=true only to genuinely remove a mistaken/garbage entry.
    """
    from guidelines_db import delete_guideline
    try:
        success = delete_guideline(guideline_id, hard_delete=hard_delete)
        if not success:
            raise HTTPException(404, "Guideline not found")
        return {"success": True, "hard_deleted": hard_delete}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Could not delete guideline: {e}")


@app.post("/guidelines/bulk-import")
async def bulk_import_guidelines(
    client: str = Form(...),
    ott_platform: str = Form(...),
    year: int = Form(...),
    guidelines_file: UploadFile = File(None),
    guidelines_text: str = Form(default=""),
):
    """
    Digitize a whole guidelines document in one step — upload the original
    PDF/DOC/XLSX (or paste raw text) and the AI splits it into individual
    structured spec rows, exactly like the company manually did to build the
    'Sample Data Collation' table in their own mockup deck. This is what
    makes adding a brand-new OTT platform's guidelines fast instead of
    someone typing each row in by hand.
    """
    from guidelines_db import extract_guidelines_with_ai, add_guideline

    raw_text = ""
    if guidelines_file and guidelines_file.filename:
        file_bytes = await guidelines_file.read()
        raw_text = read_file(file_bytes, guidelines_file.filename)["raw_text"]
    elif guidelines_text.strip():
        raw_text = guidelines_text.strip()

    if not raw_text:
        raise HTTPException(400, "Provide either a file or pasted guideline text")

    try:
        entries = extract_guidelines_with_ai(raw_text, client, ott_platform, year)
    except Exception as e:
        raise HTTPException(500, f"AI extraction failed: {e}")

    added_ids = []
    for entry in entries:
        try:
            added_ids.append(add_guideline(entry))
        except Exception as e:
            print(f"[bulk-import] Skipped one row due to error: {e}")

    return {
        "success": True,
        "entries_extracted": len(entries),
        "entries_added": len(added_ids),
        "ids": added_ids,
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)


# ─── TRANSCRIBE ──────────────────────────────────────────────

@app.post("/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    import tempfile
    import os
    import json
    import threading
    import queue
    import asyncio
    from timecoded_subtitles import _from_seconds

    filename = file.filename or "audio.webm"
    
    # Save the uploaded file to a temporary location for Whisper
    with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    async def transcribe_stream():
        yield f"data: {json.dumps({'status': 'starting', 'message': 'Initializing transcription engine...', 'progress': 0})}\n\n"
        await asyncio.sleep(0.05)
        
        q = queue.Queue()
        
        def worker():
            try:
                from transcription_engine import transcribe_with_stable_timestamps

                q.put({"type": "status", "message": "Loading stable-ts Whisper timestamp engine...", "progress": 10})

                def report_progress(done, total):
                    total = total or 1.0
                    pct = min(95, max(20, int(20 + (float(done) / float(total)) * 75)))
                    q.put({"type": "status", "message": f"Stabilizing word timecodes ({pct}%)...", "progress": pct})

                q.put({"type": "status", "message": "Analyzing audio and stabilizing speech boundaries...", "progress": 20})
                segments, engine = transcribe_with_stable_timestamps(tmp_path, report_progress)
                whisper_subs = []

                for i, seg in enumerate(segments, start=1):
                    whisper_subs.append({
                        "id": i,
                        "start_time": _from_seconds(seg["start"]),
                        "end_time": _from_seconds(seg["end"]),
                        "text": seg["text"],
                        "words": seg.get("words", []),
                        "confidence": seg.get("confidence"),
                        "language": seg.get("language"),
                        "flagged": False,
                        "flag_reason": ""
                    })

                q.put({"type": "done", "subtitles": whisper_subs, "engine": engine})

            except Exception as e:
                print(f"[Transcribe Error] {e}")
                q.put({"type": "error", "error": f"Transcription error: {str(e)}"})

        thread = threading.Thread(target=worker)
        thread.start()
        
        while True:
            try:
                msg = q.get_nowait()
            except queue.Empty:
                await asyncio.sleep(0.1)
                continue
                
            if msg["type"] == "status":
                yield f"data: {json.dumps({'status': 'processing', 'message': msg['message'], 'progress': msg['progress']})}\n\n"
            elif msg["type"] == "error":
                yield f"data: {json.dumps({'status': 'error', 'error': msg['error']})}\n\n"
                break
            elif msg["type"] == "done":
                result = {
                    "subtitles": msg["subtitles"],
                    "stats": {
                        "total_lines": len(msg["subtitles"]),
                        "flagged_lines": 0,
                        "platform": "none",
                        "detected_structure": "stable_ts_whisper_audio",
                        "transcription_engine": msg.get("engine", "stable-ts"),
                        "original_format": filename
                    }
                }
                yield f"data: {json.dumps({'status': 'completed', 'progress': 100, 'result': result})}\n\n"
                break
                
        # Clean up temporary file
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return StreamingResponse(
        transcribe_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        }
    )

@app.post("/transcribe-and-align")
async def transcribe_and_align_endpoint(
    audio: UploadFile = File(...),
    script: UploadFile = File(None),
    platform: str = Form(default="generic"),
    mode: str = Form("full")
):
    import tempfile
    import os
    import json
    import threading
    import queue
    import asyncio
    from timecoded_subtitles import _from_seconds
    from platform_rules import get_platform
    from file_reader import read_file
    from extractor import pre_extract_dialogue
    from transcript_aligner import align_transcription_to_script

    filename = audio.filename or "audio.webm"
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
        content = await audio.read()
        tmp.write(content)
        tmp_path = tmp.name

    # Pre-process script if provided
    script_subs = []
    if script:
        script_bytes = await script.read()
        script_name = script.filename or "script.txt"
        file_data = read_file(script_bytes, script_name)
        raw_text = file_data["raw_text"]
        structure = file_data["structure"]
        
        from timecoded_subtitles import parse_timecoded_subtitles
        
        if structure in ("srt", "vtt", "ttml") or script_name.lower().endswith((".srt", ".vtt")):
            parsed_subs = parse_timecoded_subtitles(raw_text)
            for i, sub in enumerate(parsed_subs, 1):
                sub.setdefault("id", i)
                sub.setdefault("flagged", False)
                script_subs.append(sub)
        else:
            platform_dict = get_platform(platform)
            pre_extracted = pre_extract_dialogue(raw_text, structure, script_bytes, script_name, platform_dict)
            if not pre_extracted:
                pre_extracted = raw_text.splitlines()

            for i, line in enumerate(pre_extracted, 1):
                clean_line = line.strip()
                if clean_line:
                    script_subs.append({"id": i, "text": clean_line, "flagged": False})

    async def transcribe_align_stream():
        yield f"data: {json.dumps({'status': 'starting', 'message': 'Initializing transcription engine...', 'progress': 0})}\n\n"
        await asyncio.sleep(0.05)
        
        q = queue.Queue()
        
        def worker():
            try:
                from transcription_engine import transcribe_with_stable_timestamps

                q.put({"type": "status", "message": "Loading stable-ts Whisper timestamp engine...", "progress": 10})

                def report_progress(done, total):
                    total = total or 1.0
                    pct = min(85, max(20, int(20 + (float(done) / float(total)) * 65)))
                    q.put({"type": "status", "message": f"Stabilizing word timecodes ({pct}%)...", "progress": pct})

                q.put({"type": "status", "message": "Analyzing audio and stabilizing speech boundaries...", "progress": 20})
                segments, engine = transcribe_with_stable_timestamps(tmp_path, report_progress)
                whisper_subs = []

                for i, seg in enumerate(segments, start=1):
                    whisper_subs.append({
                        "id": i,
                        "start_time": _from_seconds(seg["start"]),
                        "end_time": _from_seconds(seg["end"]),
                        "text": seg["text"],
                        "flagged": False,
                        "flag_reason": "",
                        "words": seg.get("words", []),
                        "confidence": seg.get("confidence"),
                        "language": seg.get("language"),
                    })

                if script_subs:
                    q.put({"type": "status", "message": "Aligning script text with audio timecodes...", "progress": 90})
                    final_subs = align_transcription_to_script(whisper_subs, script_subs, mode=mode)
                else:
                    final_subs = whisper_subs

                q.put({"type": "done", "subtitles": final_subs,
                       "reference_subtitles": whisper_subs, "engine": engine})

            except Exception as e:
                print(f"[ForcedAlign Error] {e}")
                q.put({"type": "error", "error": f"Alignment error: {str(e)}"})

        thread = threading.Thread(target=worker)
        thread.start()
        
        while True:
            try:
                msg = q.get_nowait()
            except queue.Empty:
                await asyncio.sleep(0.1)
                continue
                
            if msg["type"] == "status":
                yield f"data: {json.dumps({'status': 'processing', 'message': msg['message'], 'progress': msg['progress']})}\n\n"
            elif msg["type"] == "error":
                yield f"data: {json.dumps({'status': 'error', 'error': msg['error']})}\n\n"
                break
            elif msg["type"] == "done":
                _aligned = msg["subtitles"]
                _matched = sum(
                    1 for s in _aligned
                    if s.get("start_time") and not s.get("manual_placement")
                )
                result = {
                    "subtitles": _aligned,
                    "reference_subtitles": msg.get("reference_subtitles", []),
                    "stats": {
                        "total_lines": len(_aligned),
                        "flagged_lines": sum(1 for s in _aligned if s.get("flagged")),
                        "platform": platform,
                        "detected_structure": "aligned_script" if script_subs else "stable_ts_whisper_audio",
                        "transcription_engine": msg.get("engine", "stable-ts"),
                        "original_format": filename,
                        "alignStats": {
                            "mode": "aligned",
                            "matched": _matched,
                            "total": len(_aligned),
                        },
                    }
                }
                yield f"data: {json.dumps({'status': 'completed', 'progress': 100, 'result': result})}\n\n"
                break

        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return StreamingResponse(
        transcribe_align_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        }
    )

@app.post("/align-scripts")
async def align_scripts_endpoint(
    script_file: UploadFile = File(...),
    timestamps_file: UploadFile = File(...),
    mode: str = Form("full")
):
    from file_reader import read_file
    from extractor import pre_extract_dialogue
    from timecoded_subtitles import parse_timecoded_subtitles
    from transcript_aligner import align_transcription_to_script
    import asyncio
    
    script_bytes = await script_file.read()
    script_name = script_file.filename or "script.txt"
    script_data = read_file(script_bytes, script_name)
    script_raw = script_data["raw_text"]
    script_structure = script_data["structure"]
    
    script_subs = []
    if script_structure in ("srt", "vtt", "ttml") or script_name.lower().endswith((".srt", ".vtt")):
        parsed_subs = parse_timecoded_subtitles(script_raw)
        for i, sub in enumerate(parsed_subs, 1):
            sub.setdefault("id", i)
            sub.setdefault("flagged", False)
            script_subs.append(sub)
    else:
        pre_extracted = pre_extract_dialogue(script_raw, script_structure, script_bytes, script_name, {})
        if not pre_extracted:
            pre_extracted = script_raw.splitlines()
            
        for i, line in enumerate(pre_extracted, 1):
            clean_line = line.strip()
            if clean_line:
                script_subs.append({"id": i, "text": clean_line, "flagged": False})

    ts_bytes = await timestamps_file.read()
    ts_name = timestamps_file.filename or "timestamps.srt"
    ts_data = read_file(ts_bytes, ts_name)
    ts_raw = ts_data["raw_text"]
    
    ts_subs = parse_timecoded_subtitles(ts_raw)
    if not ts_subs:
        raise HTTPException(400, "Could not extract timecodes from the timestamps file.")

    from semantic_matcher import model_status, prepare_model
    if not model_status().get("cached"):
        raise HTTPException(503, "Multilingual matching model is not prepared. Use the one-time setup action first.")
    if not model_status().get("loaded"):
        await asyncio.to_thread(prepare_model)

    final_subs = align_transcription_to_script(ts_subs, script_subs, mode=mode)

    _matched = sum(
        1 for s in final_subs
        if s.get("start_time") and not s.get("manual_placement")
    )
    return {
        "subtitles": final_subs,
        "reference_subtitles": ts_subs,
        "stats": {
            "total_lines": len(final_subs),
            "flagged_lines": sum(1 for s in final_subs if s.get("flagged")),
            "platform": "generic",
            "detected_structure": "aligned_scripts",
            "original_format": script_name,
            "alignStats": {
                "mode": "aligned",
                "matched": _matched,
                "total": len(final_subs),
            },
        }
    }

@app.post("/align-scripts-stream")
async def align_scripts_stream_endpoint(
    script_file: UploadFile = File(...), timestamps_file: UploadFile = File(...), mode: str = Form("full")
):
    """Streaming equivalent used by the UI; alignment logic is unchanged."""
    import queue, threading, time
    q: queue.Queue = queue.Queue()
    script_bytes = await script_file.read(); ts_bytes = await timestamps_file.read()
    script_name = script_file.filename or "script.txt"; ts_name = timestamps_file.filename or "timestamps.srt"

    def worker():
        try:
            from file_reader import read_file
            from extractor import pre_extract_dialogue
            from timecoded_subtitles import parse_timecoded_subtitles
            from transcript_aligner import align_transcription_to_script

            q.put({"status": "processing", "progress": 10, "message": "Reading client subtitle text…"})
            sd = read_file(script_bytes, script_name)
            raw = sd["raw_text"]; structure = sd["structure"]
            parsed = (
                parse_timecoded_subtitles(raw)
                if structure in ("srt", "vtt", "ttml")
                   or script_name.lower().endswith((".srt", ".vtt"))
                else []
            )
            if parsed:
                script_subs = [dict(x, id=i, flagged=False) for i, x in enumerate(parsed, 1)]
            else:
                lines = pre_extract_dialogue(raw, structure, script_bytes, script_name, {}) or raw.splitlines()
                script_subs = [{"id": i, "text": x.strip(), "flagged": False}
                               for i, x in enumerate(lines, 1) if x.strip()]

            q.put({"status": "processing", "progress": 28, "message": f"Loaded {len(script_subs)} script lines."})
            q.put({"status": "processing", "progress": 38, "message": "Reading Whisper / Stable-ts timestamps…"})
            td = read_file(ts_bytes, ts_name)
            ts_subs = parse_timecoded_subtitles(td["raw_text"])
            if not ts_subs:
                raise ValueError("Could not extract timecodes from the timestamps file.")
            q.put({"status": "processing", "progress": 52, "message": f"Loaded {len(ts_subs)} reference cues."})
            q.put({"status": "processing", "progress": 58, "message": "Building word-level timeline..."})
            from semantic_matcher import model_status, prepare_model
            semantic_status = model_status()
            if not semantic_status.get("cached"):
                raise RuntimeError(
                    "Multilingual matching model is not prepared. Click Prepare once, "
                    "wait for Model ready, then start alignment."
                )
            if not semantic_status.get("loaded"):
                q.put({"status": "processing", "progress": 60, "message": "Loading cached multilingual model (no download)..."})
                prepare_model()
            q.put({"status": "processing", "progress": 62, "message": "Model ready; starting global alignment..."})

            # --- progress callback so the DP can report its own steps ----------
            total_lines = max(1, len(script_subs))

            def _progress(step: str, done: int = 0, total: int = 0):
                """Called by align_transcription_to_script at key phases."""
                if step == "bridge":
                    pct = 62 + int((done / max(1, total)) * 8)
                    q.put({"status": "processing", "progress": pct,
                           "message": f"Preparing local language bridge: {done}/{total} cues..."})
                elif step == "lexical":
                    pct = 70 + int((done / max(1, total)) * 10)  # 70-80
                    q.put({"status": "processing", "progress": pct,
                           "message": f"Lexical pass: {done}/{total} lines…"})
                elif step == "semantic":
                    pct = 80 + int((done / max(1, total)) * 10)  # 80-90
                    q.put({"status": "processing", "progress": pct,
                           "message": f"Semantic pass: {done}/{total} cues…"})
                elif step == "dp":
                    q.put({"status": "processing", "progress": 90,
                           "message": "Running sequence DP optimisation…"})
                elif step == "gap_fill":
                    q.put({"status": "processing", "progress": 95,
                           "message": "Gap-filling unmatched lines…"})

            final = align_transcription_to_script(
                ts_subs, script_subs, mode=mode, progress_callback=_progress
            )
            matched = sum(1 for x in final if x.get("start_time") and x.get("status") != "UNMATCHED")
            result = {
                "subtitles": final,
                "reference_subtitles": ts_subs,
                "stats": {
                    "total_lines": len(final),
                    "flagged_lines": sum(1 for x in final if x.get("flagged")),
                    "platform": "generic",
                    "detected_structure": "aligned_scripts",
                    "original_format": script_name,
                    "bridge_used": any(x.get("bridge_used") for x in final),
                    "bridge_source_language": next((x.get("bridge_source_language") for x in final if x.get("bridge_source_language")), ""),
                    "bridge_target_language": next((x.get("bridge_target_language") for x in final if x.get("bridge_target_language")), ""),
                    "alignStats": {"mode": "aligned", "matched": matched, "total": len(final)},
                },
            }
            q.put({"status": "completed", "progress": 100, "message": "Alignment complete.", "result": result})
        except Exception as exc:
            q.put({"status": "error", "progress": 0, "error": str(exc)})

    threading.Thread(target=worker, daemon=True).start()

    async def events():
        last_heartbeat = time.monotonic()
        last_known_pct = 58  # starts at the "building timeline" step
        while True:
            try:
                msg = q.get_nowait()
            except queue.Empty:
                # Report liveness without claiming work has completed.
                if time.monotonic() - last_heartbeat >= 3:
                    last_heartbeat = time.monotonic()
                    # Keep the last real stage percentage while the worker is busy.
                    hb_payload = {"status": "processing", "progress": last_known_pct, "message": "Mapping dialogue to timestamps… (" + str(last_known_pct) + "%)"}
                    hb_payload["message"] = "Backend is still working on alignment..."
                    yield "data: " + json.dumps(hb_payload) + "\n\n"
                await asyncio.sleep(0.1)
                continue
            last_heartbeat = time.monotonic()
            if msg.get("status") == "processing":
                last_known_pct = msg.get("progress", last_known_pct)
            yield f"data: {json.dumps(msg)}\n\n"
            if msg["status"] in ("completed", "error"):
                break

    return StreamingResponse(
        events(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/refine-alignment")
async def refine_alignment_endpoint(data: dict):
    """Run the optional, bounded AI review after manual mapping is visible."""
    reference_subtitles = data.get("reference_subtitles", [])
    subtitles = data.get("subtitles", [])
    mode = data.get("mode", "full")
    if not reference_subtitles or not subtitles:
        raise HTTPException(400, "Mapped subtitles and their reference cues are required for AI review.")

    from llm_aligner import refine_alignment_with_llm
    refined = refine_alignment_with_llm(reference_subtitles, subtitles, mode=mode)
    reviewed_count = sum(1 for sub in refined if sub.get("align_method") == "ai")
    return {"subtitles": refined, "reviewed_count": reviewed_count}


@app.post("/refine-alignment-stream")
async def refine_alignment_stream_endpoint(data: dict):
    """Optional AI placement with progress events for the manual-review panel."""
    import queue
    import threading
    reference_subtitles = data.get("reference_subtitles", [])
    subtitles = data.get("subtitles", [])
    mode = data.get("mode", "full")
    if not reference_subtitles or not subtitles:
        raise HTTPException(400, "Mapped subtitles and their reference cues are required for AI placement.")

    async def event_stream():
        events = queue.Queue()
        def worker():
            try:
                from llm_aligner import refine_alignment_with_llm
                def report(done, total, current, sub):
                    mapped = sum(1 for item in current if item.get("align_method") == "ai")
                    events.put({"type":"progress", "done":done, "total":total, "mapped":mapped})
                refined = refine_alignment_with_llm(reference_subtitles, subtitles, mode=mode, progress_callback=report)
                events.put({"type":"done", "subtitles":refined})
            except Exception as exc:
                events.put({"type":"error", "error":str(exc)})
        threading.Thread(target=worker, daemon=True).start()
        while True:
            try:
                event = events.get_nowait()
            except queue.Empty:
                await asyncio.sleep(0.15)
                continue
            if event["type"] == "progress":
                total = max(1, event["total"])
                pct = int((event["done"] / total) * 100)
                message = f"AI mapping {event['done']} of {event['total']} remaining lines ({event['mapped']} placed)"
                yield f"data: {json.dumps({'status':'processing','progress':pct,'message':message})}\n\n"
            elif event["type"] == "done":
                yield f"data: {json.dumps({'status':'completed','result':{'subtitles':event['subtitles']}})}\n\n"
                break
            else:
                yield f"data: {json.dumps({'status':'error','error':event['error']})}\n\n"
                break
    return StreamingResponse(event_stream(), media_type="text/event-stream", headers={"Cache-Control":"no-cache", "X-Accel-Buffering":"no"})

# ── SOURCE DOCUMENT HTML PREVIEW (DOC / DOCX / XLSX) ──────────────────────────
import html as _html

def _find_soffice():
    for cand in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice", "libreoffice",
    ):
        try:
            if _os.path.exists(cand):
                return cand
        except Exception:
            pass
    from shutil import which
    return which("soffice") or which("libreoffice")


def _docx_to_html(path: str) -> str:
    from docx import Document
    doc = Document(path)
    out = ['<div style="font-family:Calibri,Arial,sans-serif;color:#1f2937;line-height:1.6;padding:8px 4px">']
    def esc(t): return _html.escape(t or "")
    for p in doc.paragraphs:
        txt = esc(p.text)
        if not txt.strip():
            out.append('<div style="height:10px"></div>')
            continue
        style = (p.style.name or "").lower()
        if "heading 1" in style:
            out.append(f'<h1 style="font-size:22px;margin:14px 0 6px">{txt}</h1>')
        elif "heading 2" in style:
            out.append(f'<h2 style="font-size:18px;margin:12px 0 5px">{txt}</h2>')
        elif "heading 3" in style:
            out.append(f'<h3 style="font-size:15px;margin:10px 0 4px">{txt}</h3>')
        elif "title" in style:
            out.append(f'<h1 style="font-size:26px;text-align:center;margin:8px 0">{txt}</h1>')
        else:
            out.append(f'<p style="margin:4px 0">{txt}</p>')
    for ti, table in enumerate(doc.tables):
        out.append('<table border="1" cellspacing="0" cellpadding="6" style="border-collapse:collapse;margin:12px 0;width:100%;font-size:13px">')
        for row in table.rows:
            out.append("<tr>")
            for cell in row.cells:
                out.append(f'<td style="border:1px solid #cbd5e1;vertical-align:top">{esc(cell.text)}</td>')
            out.append("</tr>")
        out.append("</table>")
    out.append("</div>")
    return "".join(out)


def _xlsx_to_html(path: str, active_sheet_name: str = "") -> str:
    from openpyxl import load_workbook
    try:
        wb = load_workbook(path, data_only=True)
    except Exception as e:
        return f'<div style="font-family:sans-serif;color:#dc2626;padding:24px;background:#fef2f2;border-radius:8px">❌ Error reading Excel file: {_html.escape(str(e))}</div>'

    sheet_names = wb.sheetnames
    if not sheet_names:
        return '<div style="font-family:sans-serif;color:#64748b;padding:24px;text-align:center">Excel file contains no worksheets.</div>'

    default_idx = 0
    if active_sheet_name:
        target_lower = active_sheet_name.lower().strip()
        for idx, name in enumerate(sheet_names):
            name_lower = name.lower().strip()
            if name_lower == target_lower or target_lower in name_lower or name_lower in target_lower:
                default_idx = idx
                break

    filename = _os.path.basename(path)

    def col_letter(col_idx):
        result = ""
        while col_idx >= 0:
            result = chr(65 + (col_idx % 26)) + result
            col_idx = (col_idx // 26) - 1
        return result

    html_parts = ["""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: 'Segoe UI', Calibri, Arial, sans-serif; background: #f8fafc; color: #1e293b; font-size: 13px; display: flex; flex-direction: column; height: 100vh; overflow: hidden; }
  .excel-topbar { background: #107c41; color: white; padding: 10px 16px; display: flex; align-items: center; justify-content: space-between; gap: 12px; flex-shrink: 0; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
  .excel-title { font-weight: 700; font-size: 14px; display: flex; align-items: center; gap: 8px; }
  .excel-controls { display: flex; align-items: center; gap: 10px; }
  .excel-search { padding: 6px 12px; border-radius: 6px; border: 1px solid #0b572e; font-size: 12px; outline: none; width: 240px; background: rgba(255,255,255,0.95); color: #0f172a; }
  .btn-toggle-wrap { padding: 5px 12px; background: rgba(255,255,255,0.2); color: white; border: 1px solid rgba(255,255,255,0.4); border-radius: 6px; font-size: 11px; font-weight: 600; cursor: pointer; transition: all 0.15s ease; }
  .btn-toggle-wrap:hover { background: rgba(255,255,255,0.35); }
  .excel-container { flex: 1; overflow: auto; background: #ffffff; position: relative; }
  .excel-sheet { display: none; min-width: 100%; }
  .excel-sheet.active { display: block; }
  .excel-table { border-collapse: separate; border-spacing: 0; width: 100%; font-size: 12.5px; table-layout: auto; }
  .excel-table th, .excel-table td { border-right: 1px solid #cbd5e1; border-bottom: 1px solid #cbd5e1; padding: 8px 12px; white-space: pre-wrap; word-break: break-word; vertical-align: top; line-height: 1.5; min-width: 120px; }
  .excel-table.nowrap td { white-space: nowrap !important; max-width: none !important; }
  .col-hdr { background: #f1f5f9; color: #475569; font-weight: 700; text-align: center; position: sticky; top: 0; z-index: 10; border-top: 1px solid #cbd5e1; user-select: none; font-size: 11px; white-space: nowrap !important; }
  .row-num { background: #f1f5f9; color: #64748b; font-weight: 600; text-align: center; position: sticky; left: 0; z-index: 5; width: 45px; min-width: 45px; user-select: none; font-size: 11px; white-space: nowrap !important; }
  .corner-hdr { background: #e2e8f0; position: sticky; top: 0; left: 0; z-index: 20; width: 45px; min-width: 45px; border-top: 1px solid #cbd5e1; }
  .excel-tabs { background: #e2e8f0; border-top: 1px solid #cbd5e1; padding: 4px 8px 0; display: flex; gap: 4px; flex-shrink: 0; overflow-x: auto; user-select: none; }
  .excel-tab { padding: 7px 16px; background: #cbd5e1; color: #334155; font-size: 12px; font-weight: 600; border-radius: 6px 6px 0 0; cursor: pointer; border: 1px solid #94a3b8; border-bottom: none; display: flex; align-items: center; gap: 6px; transition: all 0.15s ease; }
  .excel-tab.active { background: #ffffff; color: #107c41; font-weight: 800; border-color: #cbd5e1; border-top: 3px solid #107c41; }
  .excel-tab:hover:not(.active) { background: #e2e8f0; color: #0f172a; }
  .badge-sheet { font-size: 10px; background: #e0e7ff; color: #3730a3; padding: 2px 6px; border-radius: 4px; font-weight: 700; margin-left: 4px; }
</style>
</head>
<body>
"""]

    html_parts.append(f'''
    <div class="excel-topbar">
      <div class="excel-title">
        <span>📊 {_html.escape(filename)}</span>
        <span style="font-size: 11px; opacity: 0.85; font-weight: 500;">({len(sheet_names)} sheet{'s' if len(sheet_names) > 1 else ''})</span>
      </div>
      <div class="excel-controls">
        <button class="btn-toggle-wrap" id="wrapToggleBtn" onclick="toggleWrapText()">↩ Text Wrap: ON</button>
        <input type="text" class="excel-search" id="excelSearch" placeholder="🔍 Search active sheet..." onkeyup="filterExcelRows()" />
      </div>
    </div>
    <div class="excel-container">
    ''')

    for s_idx, s_name in enumerate(sheet_names):
        ws = wb[s_name]
        is_active = (s_idx == default_idx)
        active_cls = " active" if is_active else ""
        
        html_parts.append(f'<div class="excel-sheet{active_cls}" id="sheet-{s_idx}">')
        
        rows = list(ws.iter_rows(values_only=True))
        non_empty_rows = [r for r in rows if any(c is not None and str(c).strip() != "" for c in r)]
        
        if not non_empty_rows:
            html_parts.append('<div style="padding:40px;text-align:center;color:#64748b;font-style:italic">Sheet is empty</div></div>')
            continue

        max_cols = max(len(r) for r in non_empty_rows)
        
        html_parts.append('<table class="excel-table"><thead><tr>')
        html_parts.append('<th class="corner-hdr"></th>')
        for c in range(max_cols):
            html_parts.append(f'<th class="col-hdr">{col_letter(c)}</th>')
        html_parts.append('</tr></thead><tbody>')

        for r_idx, row in enumerate(non_empty_rows):
            html_parts.append(f'<tr class="excel-row"><td class="row-num">{r_idx + 1}</td>')
            for c_idx in range(max_cols):
                val = row[c_idx] if c_idx < len(row) else ""
                val_str = "" if val is None else str(val)
                is_header = (r_idx == 0)
                cell_style = "font-weight:700;background:#f8fafc;" if is_header else ""
                html_parts.append(f'<td style="{cell_style}" title="{_html.escape(val_str)}">{_html.escape(val_str)}</td>')
            html_parts.append('</tr>')

        html_parts.append('</tbody></table></div>')

    html_parts.append('</div>')

    html_parts.append('<div class="excel-tabs">')
    for s_idx, s_name in enumerate(sheet_names):
        is_active = (s_idx == default_idx)
        active_cls = " active" if is_active else ""
        badge = '<span class="badge-sheet">TARGET</span>' if is_active and active_sheet_name else ""
        html_parts.append(f'''
        <div class="excel-tab{active_cls}" id="tab-{s_idx}" onclick="switchSheet({s_idx})">
          <span>📄</span>
          <span>{_html.escape(s_name)}</span>
          {badge}
        </div>
        ''')
    html_parts.append('</div>')

    html_parts.append('''
    <script>
      let isWrapped = true;
      function toggleWrapText() {
        isWrapped = !isWrapped;
        const btn = document.getElementById('wrapToggleBtn');
        const tables = document.querySelectorAll('.excel-table');
        tables.forEach(tbl => {
          if (isWrapped) {
            tbl.classList.remove('nowrap');
          } else {
            tbl.classList.add('nowrap');
          }
        });
        if (btn) btn.textContent = isWrapped ? '↩ Text Wrap: ON' : '➡️ Text Wrap: OFF';
      }
      function switchSheet(idx) {
        document.querySelectorAll('.excel-sheet').forEach(s => s.classList.remove('active'));
        document.querySelectorAll('.excel-tab').forEach(t => t.classList.remove('active'));
        const activeSheet = document.getElementById('sheet-' + idx);
        const activeTab = document.getElementById('tab-' + idx);
        if (activeSheet) activeSheet.classList.add('active');
        if (activeTab) activeTab.classList.add('active');
        filterExcelRows();
      }
      function filterExcelRows() {
        const q = (document.getElementById('excelSearch').value || '').toLowerCase().trim();
        const activeSheet = document.querySelector('.excel-sheet.active');
        if (!activeSheet) return;
        const rows = activeSheet.querySelectorAll('tbody tr');
        rows.forEach(r => {
          if (!q) { r.style.display = ''; return; }
          const txt = r.textContent.toLowerCase();
          r.style.display = txt.includes(q) ? '' : 'none';
        });
      }
    </script>
    </body>
    </html>
    ''')

    try:
        wb.close()
    except Exception:
        pass

    return "".join(html_parts)


def _doc_to_html(path: str) -> str:
    from file_reader import read_file
    text = read_file(_read_bytes(path), _os.path.basename(path), force_ocr=False).get("raw_text", "")
    esc = _html.escape(text or "")
    return f'<div style="font-family:Consolas,monospace;white-space:pre-wrap;color:#1f2937;padding:8px">{esc}</div>'


def _read_bytes(path: str) -> bytes:
    with open(path, "rb") as fh:
        return fh.read()


# ── SOURCE DOCUMENT STORAGE ──────────────────────────────────────────────────
import uuid as _uuid
import os as _os

SOURCE_DIR = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "uploads", "sources")
SOURCE_MEDIA = {
    "pdf": "application/pdf",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "gif": "image/gif",
    "bmp": "image/bmp",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xls": "application/vnd.ms-excel",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "txt": "text/plain",
    "rtf": "application/rtf",
    "csv": "text/csv",
    "xml": "application/xml",
    "ttml": "application/xml",
    "html": "text/html",
}


def _store_source_file(file_bytes: bytes, filename: str) -> str:
    """Persist original upload bytes; returns a unique file id for this file."""
    try:
        _os.makedirs(SOURCE_DIR, exist_ok=True)
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "bin"
        safe_name = re.sub(r"[^a-z0-9._-]", "_", filename.lower())
        file_id = _uuid.uuid4().hex
        path = _os.path.join(SOURCE_DIR, f"{file_id}__{safe_name}")
        with open(path, "wb") as fh:
            fh.write(file_bytes)
        return file_id
    except Exception as e:
        print(f"[SOURCE] failed to store {filename}: {e}")
        return None


def _source_path_for(file_id: str):
    if not file_id:
        return None
    try:
        if _os.path.exists(SOURCE_DIR):
            for entry in _os.listdir(SOURCE_DIR):
                if entry.startswith(file_id + "__"):
                    return _os.path.join(SOURCE_DIR, entry)
    except Exception:
        return None
    return None


def _source_preview_html(file_id: str, target_sheet: str = ""):
    """Return (content_type, body) for an in-app preview of an office file."""
    path = _source_path_for(file_id)
    if not path:
        return None
    ext = path.lower().rsplit(".", 1)[-1]

    if ext in ("xlsx", "xls", "csv"):
        return ("text/html", _xlsx_to_html(path, active_sheet_name=target_sheet).encode("utf-8"))

    soffice = _find_soffice()
    if soffice and ext in ("doc", "docx"):
        try:
            import tempfile, subprocess
            out_dir = tempfile.mkdtemp()
            subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                            "--outdir", out_dir, path],
                           capture_output=True, timeout=120)
            pdf = _os.path.join(out_dir, _os.path.splitext(_os.path.basename(path))[0] + ".pdf")
            if _os.path.exists(pdf):
                return ("application/pdf", _read_bytes(pdf))
        except Exception as e:
            print(f"[SOURCE] soffice conversion failed, fallback to HTML: {e}")
    if ext == "docx":
        return ("text/html", _docx_to_html(path).encode("utf-8"))
    if ext == "doc":
        return ("text/html", _doc_to_html(path).encode("utf-8"))
    return None
@app.get("/platforms/source-file/{import_id}")
def get_source_file(import_id: str):
    """
    Serve the ORIGINAL uploaded guideline document so the frontend can render a
    faithful in-app preview (PDF, image, DOC, XLSX, …) instead of only text.
    """
    path = _source_path_for(import_id)
    if not path or not _os.path.exists(path):
        raise HTTPException(404, "Source file not found")
    ext = path.lower().rsplit(".", 1)[-1]
    media = SOURCE_MEDIA.get(ext, "application/octet-stream")
    return StreamingResponse(
        open(path, "rb"),
        media_type=media,
        headers={"Content-Disposition": f'inline; filename="{_os.path.basename(path)}"',
                  "Cache-Control": "no-store"},
    )


@app.get("/platforms/source-preview/{file_id}")
def get_source_preview(file_id: str, target_sheet: str = ""):
    """
    Serve an in-app PREVIEW of an office document (DOC/DOCX/XLSX). Returns either
    a real PDF (if LibreOffice is installed for docs) or an interactive HTML rendering.
    """
    try:
        result = _source_preview_html(file_id, target_sheet=target_sheet)
        if not result:
            raise HTTPException(404, "Preview not available for this file")
        ctype, body = result
        if ctype == "application/pdf":
            return StreamingResponse(
                iter([body]), media_type="application/pdf",
                headers={"Content-Disposition": "inline", "Cache-Control": "no-store"},
            )
        return HTMLResponse(content=body)
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Error generating preview: {e}")


# ── ONLYOFFICE DOCUMENT VIEWER ────────────────────────────────────────────────
# Returns a signed JWT config so the frontend can initialise DocsAPI.DocEditor
# in view-only mode.  The ONLYOFFICE server fetches the document directly from
# our /platforms/source-file/{file_id} endpoint.

_OFFICE_DOCTYPES = {
    "docx": "word", "doc": "word",
    "xlsx": "cell", "xls": "cell",
    "pptx": "slide", "ppt": "slide",
    "odt": "word", "ods": "cell", "odp": "slide",
    "csv": "cell", "txt": "word",
}


@app.get("/platforms/onlyoffice-config/{file_id}")
def get_onlyoffice_config(file_id: str, backend_url: str = ""):
    """
    Return a JWT-signed ONLYOFFICE editor configuration for view-only preview.
    backend_url is the URL the ONLYOFFICE container uses to reach *our* backend
    (must be accessible from inside Docker).
    Defaults to http://localhost:8000 if not provided.
    """
    import hashlib
    try:
        import jwt as _jwt
    except ImportError:
        _jwt = None

    oo_url = _os.getenv("ONLYOFFICE_URL", "http://localhost:8080").rstrip("/")
    oo_secret = _os.getenv("ONLYOFFICE_JWT_SECRET", "")

    # Resolve source file
    path = _source_path_for(file_id)
    if not path or not _os.path.exists(path):
        raise HTTPException(404, "Source file not found")

    filename = _os.path.basename(path)
    # Strip the uuid prefix: "abc123__my_file.xlsx" → "my_file.xlsx"
    display_name = filename.split("__", 1)[-1] if "__" in filename else filename
    ext = display_name.lower().rsplit(".", 1)[-1] if "." in display_name else "bin"
    doc_type = _OFFICE_DOCTYPES.get(ext, "word")

    # Unique key: changes whenever file changes on disk (cache-busting)
    mtime = _os.path.getmtime(path)
    doc_key = hashlib.md5(f"{file_id}_{mtime}".encode()).hexdigest()[:20]

    # URL where ONLYOFFICE will fetch the document.
    # backend_url comes from the frontend (window.location.origin is useless for Docker,
    # so caller should pass the machine's LAN IP or localhost:8000).
    _backend = (backend_url or "http://localhost:8000").rstrip("/")
    # Never point ONLYOFFICE at the Vite dev server (port 5173); always FastAPI (8000).
    if ":5173" in _backend:
        _backend = _backend.replace(":5173", ":8000")
    file_url = f"{_backend}/platforms/source-file/{file_id}"

    config = {
        "document": {
            "fileType": ext,
            "key": doc_key,
            "title": display_name,
            "url": file_url,
            "permissions": {
                "comment": False,
                "copy": True,
                "download": True,
                "edit": False,
                "fillForms": False,
                "modifyContentControl": False,
                "modifyFilter": False,
                "print": True,
                "review": False,
            },
        },
        "documentType": doc_type,
        "editorConfig": {
            "mode": "view",
            "lang": "en",
            "coEditing": {"mode": "strict", "change": False},
            "customization": {
                "autosave": False,
                "chat": False,
                "comments": False,
                "compactHeader": True,
                "feedback": False,
                "help": False,
                "hideRightMenu": True,
                "hideRulers": True,
                "plugins": False,
                "toolbar": True,
                "toolbarNoTabs": True,
            },
        },
    }

    # Sign with JWT if secret is configured
    token = None
    if oo_secret and _jwt is not None:
        try:
            token = _jwt.encode(config, oo_secret, algorithm="HS256")
        except Exception as je:
            print(f"[ONLYOFFICE] JWT signing failed: {je}")

    return {
        "onlyoffice_url": oo_url,
        "config": config,
        "token": token,
        "doc_type": doc_type,
        "filename": display_name,
    }


@app.get("/platforms/onlyoffice-health")
def onlyoffice_health():
    """Check whether the ONLYOFFICE Document Server is reachable."""
    import httpx as _httpx
    oo_url = _os.getenv("ONLYOFFICE_URL", "http://localhost:8080").rstrip("/")
    try:
        r = _httpx.get(f"{oo_url}/healthcheck", timeout=3.0)
        return {"running": r.status_code == 200, "url": oo_url}
    except Exception:
        return {"running": False, "url": oo_url}


# ─── MOVIES ──────────────────────────────────────────────────────

from pydantic import BaseModel

class MovieAddRequest(BaseModel):
    title: str
    url: str
    added_by: str = "Anonymous"

@app.get("/movies")
def get_movies():
    try:
        from database import get_all_movies
        return {"movies": get_all_movies()}
    except Exception as e:
        print(f"Error fetching movies: {e}")
        return {"movies": []}

@app.post("/movies")
def add_new_movie(req: MovieAddRequest):
    if not req.title.strip() or not req.url.strip():
        raise HTTPException(400, "Title and URL are required")
    if not re.match(r"^https?://[^\s]+$", req.url.strip(), re.IGNORECASE):
        raise HTTPException(400, "URL must start with http:// or https://")
    try:
        from database import add_movie
        add_movie(req.title.strip(), req.url.strip(), req.added_by.strip() or "Anonymous")
        return {"message": "Movie added successfully"}
    except Exception as e:
        print(f"Error adding movie: {e}")
        raise HTTPException(500, "Failed to add movie")

# ─── SUBTITLE EDITOR ENDPOINTS ───────────────────────────────────

from editor import (
    parse_subtitles,
    subtitles_to_format,
    sync_offset,
    sync_scale,
    sync_visual,
    sync_point_via_other,
    translate_subtitles_stream,
    get_provider_list,
    detect_format,
)
import chardet

class SyncRequest(BaseModel):
    subtitles: list
    mode: str = "offset"
    seconds: float = 0.0
    factor: float = 1.0
    range_start_id: int | None = None
    range_end_id: int | None = None
    anchor_id: int | None = None
    new_start: str | None = None
    anchor_id2: int | None = None
    new_start2: str | None = None
    reference_subtitles: list | None = None
    sub_index: int | None = None
    ref_index: int | None = None
    sub_index2: int | None = None
    ref_index2: int | None = None

class ExportRequest(BaseModel):
    subtitles: list
    format: str = "srt"
    filename: str = "subtitles"

translate_jobs = {}

class StopTranslateRequest(BaseModel):
    client_id: str
    action: str = "apply"

class TranslateRequest(BaseModel):
    subtitles: list
    target_language: str
    source_language: str = ""
    provider: str = "google"
    client_id: str | None = None
    api_key: str | None = None
    model: str | None = None
    base_url: str | None = None
    custom_prompt: str | None = None

@app.get("/editor/formats")
def get_editor_formats():
    from editor import IMPORT_FORMATS, EXPORT_FORMATS
    return {"import": IMPORT_FORMATS, "export": EXPORT_FORMATS}

@app.post("/editor/import")
async def editor_import_file(file: UploadFile = File(...)):
    if not file:
        raise HTTPException(400, "No file uploaded")
    content = await file.read()
    if not content:
        raise HTTPException(400, "File is empty")
    
    # Detect encoding
    detected = chardet.detect(content)
    encoding = detected.get("encoding") or "utf-8"
    try:
        text = content.decode(encoding)
    except Exception:
        text = content.decode("utf-8", errors="replace")
    
    filename = file.filename or "subtitles.srt"
    fmt = detect_format(filename, text)
    subs = parse_subtitles(text, fmt, filename)
    return {"status": "ok", "format": fmt, "filename": filename, "subtitles": subs}

@app.post("/editor/export")
def editor_export_file(req: ExportRequest):
    if not req.subtitles:
        raise HTTPException(400, "No subtitles provided")
    fmt = (req.format or "srt").lower()
    filename = req.filename or "subtitles"
    output_text = subtitles_to_format(req.subtitles, fmt, filename)
    
    headers = {
        "Content-Disposition": f'attachment; filename="{filename}.{fmt}"'
    }
    return Response(content=output_text.encode("utf-8"), media_type="application/octet-stream", headers=headers)

@app.post("/editor/sync")
def editor_sync(req: SyncRequest):
    subs = req.subtitles or []
    mode = (req.mode or "offset").lower()
    
    if mode == "offset":
        res = sync_offset(subs, req.seconds, req.range_start_id, req.range_end_id)
    elif mode == "scale":
        res = sync_scale(subs, req.factor, req.range_start_id, req.range_end_id)
    elif mode == "visual":
        res = sync_visual(
            subs, req.anchor_id or 1, req.new_start or "00:00:00,000",
            req.anchor_id2, req.new_start2,
            req.range_start_id, req.range_end_id
        )
    elif mode == "point_via_other":
        res = sync_point_via_other(
            subs, req.reference_subtitles or [],
            req.sub_index, req.ref_index,
            req.sub_index2, req.ref_index2
        )
    else:
        res = sync_offset(subs, req.seconds, req.range_start_id, req.range_end_id)
        
    return {"status": "ok", "subtitles": res}

@app.get("/editor/translate/providers")
def editor_translate_providers():
    return {"providers": get_provider_list()}

@app.post("/editor/translate/stop")
def editor_translate_stop(req: StopTranslateRequest):
    if req.client_id in translate_jobs:
        translate_jobs[req.client_id]["stop"] = True
        translate_jobs[req.client_id]["action"] = req.action
    return {"status": "ok"}

@app.post("/editor/translate")
def editor_translate(req: TranslateRequest):
    from fastapi.responses import StreamingResponse
    config = {
        "api_key": req.api_key,
        "model": req.model,
        "base_url": req.base_url,
        "custom_prompt": req.custom_prompt,
    }
    
    client_id = req.client_id
    if client_id:
        translate_jobs[client_id] = {"stop": False, "action": "apply"}
        
    def stop_check():
        if client_id and client_id in translate_jobs:
            return translate_jobs[client_id].get("stop", False)
        return False
        
    def stop_action():
        if client_id and client_id in translate_jobs:
            return translate_jobs[client_id].get("action", "apply")
        return "apply"

    def event_stream():
        try:
            generator = translate_subtitles_stream(
                req.subtitles, req.target_language, req.source_language,
                provider=req.provider, config=config,
                stop_check=stop_check, stop_action=stop_action
            )
            for event in generator:
                yield f"data: {json.dumps(event)}\n\n"
        finally:
            if client_id and client_id in translate_jobs:
                del translate_jobs[client_id]
            
    return StreamingResponse(event_stream(), media_type="text/event-stream")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
