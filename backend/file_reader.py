# file_reader.py — Reads ANY subtitle/script file format
# Handles: DOC, DOCX, PDF, SRT, VTT, XML, TTML, RTF, TXT, XLSX, XLS, CSV, PMW
# Preserves structure for AI to understand context

import io
import re
import chardet


class UnsupportedPMWError(ValueError):
    pass
_PDF_METADATA_LINE = re.compile(r'^\s*(?:===|Original\s+title:|Translated\s+title:|Language:|File:|Printed\s+on\s+.+\bPage\s+\d+\s*$)', re.IGNORECASE)
_PDF_TIMECODE = re.compile(r'\b\d{2}:\d{2}:\d{2}[.:;]\d{2}\b')
# Footage timecode: FEET.FRAMES (e.g. 56.11, 83.03) — used in Hollywood CCSL documents
# Frames are 0-15 for 35mm 24fps. Feet are typically 2-4 digits.
_FOOTAGE_TC = re.compile(r'\b(\d{1,4})\.(\d{2})\b')


def _footage_to_seconds(feet: int, frames: int, fps: int = 24, frames_per_foot: int = 16) -> float:
    """Convert 35mm footage timecode to seconds. Default: 24fps, 16 frames/foot."""
    return (feet * frames_per_foot + frames) / fps


def _footage_to_hmsf(footage_str: str) -> str:
    """
    Convert a footage timecode string like '56.11' to HH:MM:SS:FF format.
    Returns empty string if not a valid footage timecode.
    """
    m = _FOOTAGE_TC.fullmatch(footage_str.strip())
    if not m:
        return ''
    feet, frames = int(m.group(1)), int(m.group(2))
    if frames > 15:  # not a valid 35mm frame count — probably a decimal number
        return ''
    total_frames = feet * 16 + frames
    total_seconds = total_frames / 24
    h = int(total_seconds // 3600)
    m2 = int((total_seconds % 3600) // 60)
    s = int(total_seconds % 60)
    f = total_frames % 24
    return f'{h:02d}:{m2:02d}:{s:02d}:{f:02d}'

def _pdf_text_needs_outline_ocr(text: str) -> bool:
    timecode_lines = word_count = 0
    for line in text.splitlines():
        line = line.strip()
        if not line or _PDF_METADATA_LINE.match(line):
            continue
        if _PDF_TIMECODE.search(line):
            timecode_lines += 1
            line = _PDF_TIMECODE.sub(' ', line)
            line = re.sub(r'\b\d+(?:[.:]\d+)*\b', ' ', line)
        word_count += len(re.findall(r"[A-Za-z][A-Za-z'-]{1,}", line))
    return timecode_lines >= 20 and word_count < max(20, timecode_lines // 5)


def _merge_pdf_timecodes_with_ocr(file_bytes: bytes, raw_text: str) -> str:
    from ocr_reader import render_pdf_pages_as_images, ocr_image_bytes
    import re
    from concurrent.futures import ThreadPoolExecutor
    
    pages_text = re.split(r'===\s*PAGE\s*\d+\s*===', raw_text)
    if pages_text and not pages_text[0].strip():
        pages_text.pop(0)
        
    images = render_pdf_pages_as_images(file_bytes)
    
    def process_page(img):
        try:
            return ocr_image_bytes(img).strip()
        except Exception as e:
            return f'[OCR Failed: {e} - Ensure Tesseract OCR is installed]'

    with ThreadPoolExecutor(max_workers=8) as executor:
        ocr_results = list(executor.map(process_page, images))

    merged = []
    for i, ocr_result in enumerate(ocr_results):
        page_num = i + 1
        merged.append(f'=== PAGE {page_num} ===')
        
        if i < len(pages_text):
            merged.append('--- ACCURATE TIMECODES ---')
            merged.append(pages_text[i].strip())
            
        merged.append('--- OCR SUBTITLES (Match with timecodes above) ---')
        merged.append(ocr_result)
        merged.append('')
        
    return '\n'.join(merged)

def _full_pdf_ocr(file_bytes: bytes) -> str:
    from ocr_reader import render_pdf_pages_as_images, ocr_image_bytes
    from concurrent.futures import ThreadPoolExecutor
    
    images = render_pdf_pages_as_images(file_bytes)
    def process_page(img):
        try:
            return ocr_image_bytes(img).strip()
        except Exception as e:
            return f'[OCR Failed: {e}]'

    with ThreadPoolExecutor(max_workers=8) as executor:
        ocr_results = list(executor.map(process_page, images))
        
    merged = []
    for i, ocr_result in enumerate(ocr_results):
        merged.append(f'=== PAGE {i+1} ===')
        merged.append(ocr_result)
        merged.append('')
    return '\n'.join(merged)

def read_file(file_bytes: bytes, filename: str, force_ocr: bool = False) -> dict:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "txt"

    # ── Magic-byte sniffing: detect misnamed files ──────────────────────────
    # A file saved as .pdf but actually a DOCX/ZIP will start with PK (0x50 0x4B)
    # A file saved as .docx but actually PDF will start with %PDF
    if len(file_bytes) >= 4:
        magic = file_bytes[:4]
        if ext == "pdf" and magic[:2] == b'PK':
            # ZIP magic — try as DOCX first
            ext = "docx"
        elif ext in ("docx", "doc") and magic[:4] == b'%PDF':
            ext = "pdf"

    readers = {
        "docx": read_docx, "doc": read_legacy_doc,
        "pdf": read_pdf,
        "srt": read_srt,
        "vtt": read_vtt, "webvtt": read_vtt,
        "xml": read_xml, "ttml": read_xml, "dfxp": read_xml,
        "rtf": read_rtf,
        "xlsx": read_excel, "xls": read_excel,
        "csv": read_csv,
        "txt": read_plain,
        "json": read_json,
        "pmw": read_pmw,
    }

    reader = readers.get(ext, read_plain)

    # Standalone image upload (a pasted spec screenshot, a scanned page saved
    # as a .png/.jpg). There is no "text" to read — go straight to OCR so the
    # image is never passed raw to the LLM (which would reject it with an
    # "image input not supported" error). The result is plain text.
    if ext in ("png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff"):
        from ocr_reader import ocr_image_bytes
        try:
            raw_text = ocr_image_bytes(file_bytes).strip()
        except Exception as e:
            print(f"[file_reader] OCR failed for image '{filename}': {e}")
            raw_text = ""
    elif force_ocr and ext == 'pdf':
        try:
            raw_text = _full_pdf_ocr(file_bytes)
        except Exception:
            raw_text = decode_bytes(file_bytes)
    else:
        try:
            raw_text = reader(file_bytes)
        except UnsupportedPMWError:
            raise
        except Exception:
            raw_text = decode_bytes(file_bytes)

    # ── Garbled-output guard ─────────────────────────────────────────────────
    # If the extracted text has too many non-printable / non-ASCII chars,
    # or if it contains obvious DOCX zip markers, the PDF was probably
    # image-only or corrupted/misnamed. Try the other reader.
    if raw_text and ext == "pdf":
        sample = raw_text[:500].replace('\n', '').replace('\r', '')
        printable = sum(1 for c in sample if c.isprintable())
        # If it looks like binary garbage, or if it's literally a DOCX
        if (len(sample) > 0 and printable / len(sample) < 0.60) or "word/document.xml" in raw_text[:1000]:
            # Likely binary garbage or a disguised DOCX — try DOCX as fallback
            try:
                alt = read_docx(file_bytes)
                if alt and alt.strip():
                    raw_text = alt
                    ext = "docx"
            except Exception:
                pass

    if not raw_text or not raw_text.strip():
        raw_text = decode_bytes(file_bytes)

    # ── UNIVERSAL OCR PASS ─────────────────────────────────────────────────────
    # Guideline / script documents in ANY format (PDF, DOCX, XLSX, RTF, DOC, XLS)
    # routinely mix real text with embedded screenshots (e.g. a pasted table of
    # rules, a scanned spec page, a logo-bearing instruction). A native text
    # reader alone silently drops that image content. Per production requirement,
    # OCR MUST run on every page / every embedded image of the document so no
    # rule that lives inside an image is ever missed. Tesseract is required; if
    # it isn't installed we degrade gracefully and keep the native text.
    raw_text = _run_universal_ocr(ext, file_bytes, raw_text)

    structure = detect_structure(raw_text, ext)

    return {
        "raw_text": raw_text,
        "structure": structure,
        "filename": filename,
        "format": ext.upper()
    }



def detect_structure(text: str, ext: str) -> str:
    if ext == "srt" or re.search(r'\d+\n\d{2}:\d{2}:\d{2},\d{3}\s+-->', text):
        return "srt_format"
    if ext in ["vtt", "webvtt"] or text.startswith("WEBVTT"):
        return "vtt_format"
    if ext in ["xml", "ttml", "dfxp"] or "<tt " in text or "<body>" in text:
        return "xml_ttml"
    # Footage-timecode CCSL — Hollywood format: FEET.FRAMES (e.g. 56.11, 83.03)
    # These documents have columns: SC# | FOOTAGE/SHOT DESCRIPTION/DIALOGUE | TITLE# | TITLE | START | FINISH | TOTAL
    if re.search(r'COMBINED CONTINUITY', text, re.IGNORECASE) or re.search(r'SPOTTING LIST TITLES', text, re.IGNORECASE):
        # Check for footage timecodes (no HH:MM:SS in the document)
        footage_hits = len(_FOOTAGE_TC.findall(text))
        hms_hits = len(re.findall(r'\d{2}:\d{2}:\d{2}', text))
        if footage_hits > hms_hits:  # predominantly footage format
            return 'ccsl_footage'

    # CCSL / Spotting List detection — check for header with TimeIn/TimeOut/Titles columns
    # This fires both for standard CCSL PDFs and for spatial-parser output
    if re.search(r'(TimeIn|Time\s+In)\s*\|.*(TimeOut|Time\s+Out)\s*\|.*(Titles?|Dialogue)', text, re.IGNORECASE):
        return "ccsl_double_dialogue"
    if "COMBINED CONTINUITY" in text or "CCSL" in text.upper() or "SPOTTING LIST" in text.upper():
        return "ccsl_double_dialogue"

    # Table with timecodes — includes spatial parser output. Matches BOTH
    # frame-accurate HH:MM:SS:FF (e.g. CCSL spotting lists) AND second-only
    # HH:MM:SS (e.g. documentary-style "TIME CODE | VISUALS | AUDIO" tables
    # like the CAR SOS/Food Factory format, which has no frame component at
    # all)
    if any(
        '|' in line and re.search(r'\d{2}:\d{2}:\d{2}([:;]\d{2})?', line)
        for line in text.splitlines()
    ):
        return "table_with_timecodes"
    if "=== TABLE" in text and re.search(r'\d{2}:\d{2}:\d{2}', text):
        return "table_with_timecodes"
    if re.search(r'\d{2}:\d{2}:\d{2}[:;]\d{2}', text) and re.search(r'(INT\.|EXT\.|OS\)|VO\))', text):
        return "paragraph_without_table"
    if re.search(r'(INT\.|EXT\.|TEASER|ACT ONE|SCENE)', text, re.IGNORECASE):
        return "paragraph_with_speaker"
    if ext in ["xlsx", "xls", "csv"] or "=== SHEET" in text:
        return "excel_spotting_list"
    
    lines = [l for l in text.splitlines() if l.strip()]
    if len(lines) > 5 and all(len(l) < 120 for l in lines[:20]):
        return "plain_script"
    return "unknown"



def decode_bytes(b: bytes) -> str:
    detected = chardet.detect(b)
    enc = detected.get("encoding") or "utf-8"
    try:
        return b.decode(enc, errors="ignore")
    except:
        return b.decode("utf-8", errors="ignore")


def _tesseract_available() -> bool:
    """Best-effort check that Tesseract OCR is installed and importable."""
    try:
        import pytesseract  # noqa: F401
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        pytesseract.get_tesseract_version()
        return True
    except Exception as e:
        print(f"[file_reader] Tesseract OCR not available — skipping universal OCR: {e}")
        return False


def _run_universal_ocr(ext: str, file_bytes: bytes, native_text: str) -> str:
    """
    Production-grade OCR pass that runs on EVERY page / embedded image of the
    document, for every supported guideline format. Guideline docs always mix
    real text with screenshots (pasted rule tables, scanned spec pages, etc.)
    so OCR must not be conditional on whether native text 'looked empty'.

    Strategy per format:
      - pdf  : render EVERY page to an image and OCR it (catches both fully
               scanned pages and embedded screenshots alike). If the PDF also
               has real text, the OCR result is appended alongside it so the
               AI sees both.
      - docx : OCR every embedded image (word/media) — pasted tables/screenshots.
      - xlsx : OCR every embedded image (xl/media) — pasted rule sheets.
      - rtf  : OCR every \\pict-embedded image.
      - doc/xls (legacy binary): OCR via raw byte-carved images where possible.
      - txt/csv/json: nothing to OCR (plain text, no images).

    Returns the original text plus any OCR content, never less than what the
    native reader produced.
    """
    if not _tesseract_available():
        return native_text

    ocr_text = ""
    try:
        if ext == "pdf":
            from ocr_reader import render_pdf_pages_as_images, ocr_image_bytes
            images = render_pdf_pages_as_images(file_bytes)
            if images:
                sections = []
                for i, img in enumerate(images):
                    page_text = ocr_image_bytes(img).strip()
                    if page_text:
                        sections.append(f"=== OCR PAGE {i+1} ===\n{page_text}")
                ocr_text = "\n\n".join(sections)
        elif ext == "docx":
            from ocr_reader import ocr_fallback_for_docx
            ocr_text = ocr_fallback_for_docx(file_bytes, len(native_text))
        elif ext in ("xlsx", "xls"):
            from ocr_reader import ocr_fallback_for_xlsx, ocr_fallback_for_xls
            ocr_text = (ocr_fallback_for_xlsx(file_bytes, len(native_text))
                        if ext == "xlsx"
                        else ocr_fallback_for_xls(file_bytes, len(native_text)))
        elif ext == "rtf":
            from ocr_reader import ocr_fallback_for_rtf
            ocr_text = ocr_fallback_for_rtf(file_bytes, len(native_text))
        elif ext == "doc":
            from ocr_reader import extract_images_from_xls, ocr_image_bytes
            # Legacy .doc is an OLE container; reuse the binary image carver.
            for img in extract_images_from_xls(file_bytes):
                t = ocr_image_bytes(img).strip()
                if t:
                    ocr_text = (ocr_text + "\n" + t).strip()
    except Exception as e:
        print(f"[file_reader] Universal OCR pass failed for '{ext}': {e}")
        return native_text

    if not ocr_text:
        return native_text

    # If native text already exists, keep it AND append OCR so the AI sees both.
    # If native text was empty, OCR alone is the content.
    if native_text and native_text.strip():
        return native_text + "\n\n=== OCR EXTRACTED CONTENT ===\n" + ocr_text
    return ocr_text


def _useful_pmw_text(text: str) -> bool:
    if not text:
        return False
    words = re.findall(r"[^\W\d_]{2,}", text, re.UNICODE)
    timecodes = re.findall(r"\b\d{1,2}[:.]\d{2}[:.]\d{2}(?:[,.:;]\d{1,3})?\b", text)
    printable = sum(c.isprintable() or c in "\r\n\t" for c in text)
    return printable / max(len(text), 1) >= 0.75 and (len(words) >= 4 or (timecodes and len(words) >= 1))


def read_pmw(file_bytes: bytes) -> str:
    """Read text/XML, ZIP, SQLite, and embedded-text Poliscript PMW variants."""
    import zipfile
    candidates = []
    if file_bytes.startswith(b"PK"):
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
                for info in archive.infolist():
                    name = info.filename.lower()
                    if not info.is_dir() and info.file_size <= 20 * 1024 * 1024 and name.endswith((".xml", ".json", ".txt", ".srt", ".vtt", ".csv", ".pmw")):
                        candidates.append(decode_bytes(archive.read(info)))
        except (zipfile.BadZipFile, RuntimeError):
            pass
    if file_bytes.startswith(b"SQLite format 3\x00"):
        import os, sqlite3, tempfile
        path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".pmw", delete=False) as tmp:
                tmp.write(file_bytes); path = tmp.name
            db = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
            rows = []
            for (table,) in db.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall():
                safe_table = table.replace('"', '""')
                try:
                    for row in db.execute(f'SELECT * FROM "{safe_table}"'):
                        values = [str(v).strip() for v in row if isinstance(v, str) and v.strip()]
                        if values: rows.append(" | ".join(values))
                except sqlite3.DatabaseError:
                    continue
            db.close(); candidates.append("\n".join(rows))
        except (OSError, sqlite3.DatabaseError):
            pass
        finally:
            if path:
                try: os.unlink(path)
                except OSError: pass
    candidates.extend([
        decode_bytes(file_bytes),
        "\n".join(m.decode("ascii", errors="ignore") for m in re.findall(rb"[\x20-\x7e]{4,}", file_bytes)),
        "\n".join(m.decode("utf-16-le", errors="ignore") for m in re.findall(rb"(?:[\x20-\x7e]\x00){4,}", file_bytes)),
        "\n".join(m.decode("utf-16-be", errors="ignore") for m in re.findall(rb"(?:\x00[\x20-\x7e]){4,}", file_bytes)),
    ])
    useful = [t.replace("\x00", "") for t in candidates if _useful_pmw_text(t.replace("\x00", ""))]
    if not useful:
        raise UnsupportedPMWError("This PMW file uses a closed, proprietary binary format that cannot be directly decoded. Please export the file as an SRT, VTT, or TXT from your GTS/Iyuno software and upload that instead.")
    def score(text):
        return len(re.findall(r"[^\W\d_]{2,}", text, re.UNICODE)) + 8 * len(re.findall(r"\d{1,2}:\d{2}:\d{2}", text))
    return max(useful, key=score)


def read_docx(file_bytes: bytes) -> str:
    sections = []
    extracted_text = ""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paras:
            sections.append("=== PARAGRAPHS ===")
            sections.extend(paras)

        # Some production scripts use one large cell per column and align
        # timecodes/audio with blank paragraphs (instead of using one Word row
        # per cue). ``cell.text`` flattens those columns independently, which
        # destroys their alignment. Keep the last recognised column layout so
        # continuation tables without a repeated header also work.
        # Matches a complete timecode token (colon or dot separated)
        tc_re = re.compile(r"^\d{1,2}[:.][\d]{2}[:.][\d]{2}(?:[,.:;]\d{1,3})?$")
        # Matches a timecode anywhere in a line (for stacked-cell detection)
        tc_line_re = re.compile(r"\b\d{1,2}[:.][\d]{2}[:.][\d]{2}(?:[,.:;]\d{1,3})?\b")

        # Keywords that identify a cell as visual/actuality description (NOT spoken dialogue)
        _VISUAL_KEYWORDS = re.compile(
            r'^\s*(?:SCENE\s*\d+|Actuality[:\s]|ACTUALITY|VISUALS?:|GRAPHICS?\s*(?:ON\s*SCREEN)?:|'
            r'SHOT\s+DESCRIPTION|INT\.|EXT\.|CU\s|WS\s|MS\s|'
            r'pan(?:ning)?|conveyor|factory\s+floor|worker\s+(?:taking|adding|pour))',
            re.IGNORECASE
        )

        def header_role(value: str):
            """Map a column header text to its role: time / dialogue / speaker / skip / None."""
            key = re.sub(r"[^a-z]", "", value.lower())
            # Timecode columns
            if (key in {"timecode", "timecodes", "time", "tc", "timein",
                        "starttime", "timestart", "in", "intime"}
                    or key.startswith("timecode")):
                return "time"
            # Dialogue / narration columns — the spoken content that becomes subtitles
            if key in {"audio", "dialogue", "dialog", "narration", "speech",
                       "subtitle", "subtitles", "text", "spoken", "english",
                       "transcript", "translation", "titles", "title",
                       "dubtitle", "captioning"}:
                return "dialogue"
            # Visual / scene / graphics columns — NOT dialogue, skip entirely
            if key in {"visuals", "visual", "graphicsonscreen", "graphiconscreen",
                       "graphics", "scene", "scenedescription", "scenedescrip",
                       "description", "actuality", "action", "onscreen",
                       "onscreentext", "video", "notes", "comment", "comments",
                       "shot", "shotdescription", "picture", "imagery"}:
                return "skip"
            # Speaker / character columns
            if key in {"character", "characters", "speaker", "char",
                       "name", "role", "cast", "who"}:
                return "speaker"
            return None

        def _cell_is_visual(cell_text: str) -> bool:
            """Heuristic: does this cell contain visual/scene descriptions, not spoken words?"""
            lines = [l.strip() for l in cell_text.splitlines() if l.strip()]
            if not lines:
                return False
            hits = sum(1 for l in lines if _VISUAL_KEYWORDS.search(l))
            return hits / len(lines) >= 0.35

        def _cell_is_timecodes(cell_text: str) -> bool:
            """True when the majority of non-blank lines in the cell are timecodes."""
            lines = [l.strip() for l in cell_text.splitlines() if l.strip()]
            if not lines:
                return False
            tc_hits = sum(1 for l in lines if tc_line_re.search(l))
            return tc_hits / len(lines) >= 0.50

        def _detect_bmsub_stacked(table) -> dict | None:
            """
            Detect FoodFactory / HouseHunters / BMSub documentary format.

            These tables have very few rows (often 1), where each row covers
            one entire scene:
              col 0  →  stacked timecodes  (01:02:31\n01:02:35\n...)
              col 1  →  visual descriptions (SCENE 5\nActuality:\n...)
              col 2  →  narration / dialogue text

            Returns {"time": col_idx, "dialogue": col_idx} when detected,
            None otherwise.
            """
            if not table.rows or len(table.columns) < 2:
                return None
            # Only activate for small tables (large tables use header detection)
            if len(table.rows) > 15:
                return None
            # Sample up to 3 data rows
            sample = list(table.rows[:3])
            votes = {}   # col_idx → {"tc": int, "visual": int, "dlg": int}
            for row in sample:
                # Deduplicate merged-cell text
                cells_raw = [c.text.strip() for c in row.cells]
                cells = []
                for c in cells_raw:
                    if not cells or c != cells[-1]:
                        cells.append(c)
                for ci, cell_text in enumerate(cells):
                    if ci not in votes:
                        votes[ci] = {"tc": 0, "visual": 0, "dlg": 0}
                    if _cell_is_timecodes(cell_text):
                        votes[ci]["tc"] += 1
                    elif _cell_is_visual(cell_text):
                        votes[ci]["visual"] += 1
                    elif cell_text.strip():
                        votes[ci]["dlg"] += 1
            if not votes:
                return None
            tc_col  = max(votes, key=lambda c: votes[c]["tc"])
            dlg_col = max(
                (c for c in votes if c != tc_col),
                key=lambda c: votes[c]["dlg"],
                default=None
            )
            # Require at least some timecode evidence
            if votes[tc_col]["tc"] == 0 or dlg_col is None:
                return None
            return {"time": tc_col, "dialogue": dlg_col}

        for i, table in enumerate(doc.tables):
            sections.append(f"\n=== TABLE {i+1} ({len(table.rows)} rows x {len(table.columns)} cols) ===")
            first_values = [c.text.strip() for c in table.rows[0].cells] if table.rows else []

            # ── Step 1: Try recognised header row ─────────────────────────────
            roles = {}
            for idx, value in enumerate(first_values):
                role = header_role(value)
                if role and role not in roles:   # first match wins per role type
                    roles[role] = idx
            first_is_header = "time" in roles and "dialogue" in roles

            # ── Step 2: FBoyIsland / dialogue-list no-header detection ─────────
            # First cell is a single bare timecode → infer columns from position
            if not first_is_header and table.rows:
                for ci, val in enumerate(first_values):
                    if tc_re.match(val.strip()):
                        roles["time"] = ci
                        for di in range(len(first_values) - 1, -1, -1):
                            if di != ci and not tc_re.match(first_values[di]):
                                if not _cell_is_visual(first_values[di]):
                                    roles["dialogue"] = di
                                    break
                        if "dialogue" in roles:
                            first_is_header = True
                            roles["no_header"] = True
                        break

            # ── Step 3: BMSub stacked single-row format (FoodFactory etc.) ─────
            # Cells contain stacked timecodes + visual descriptions + narration.
            # Standard header / single-TC detection both fail here; use heuristics.
            if not first_is_header:
                bmsub = _detect_bmsub_stacked(table)
                if bmsub:
                    roles.update(bmsub)
                    first_is_header = True
                    roles["no_header"] = True   # row 0 is data, not a header

            # ── Step 4: Emit table content ────────────────────────────────────
            if first_is_header:
                time_col = roles["time"]
                dlg_col  = roles["dialogue"]
                sections.append("TIMECODE | DIALOGUE")
                cur_tc  = [None]
                cur_dlg = [[]]

                def _flush_cue():
                    if cur_tc[0] and cur_dlg[0]:
                        sections.append(f"{cur_tc[0]} | {chr(10).join(cur_dlg[0])}")

                start_row = 0 if roles.get("no_header") else 1
                for row in table.rows[start_row:]:
                    # Deduplicate merged cells (python-docx repeats merged cell text)
                    cells_raw = [c.text.strip() for c in row.cells]
                    cells = []
                    for c in cells_raw:
                        if not cells or c != cells[-1]:
                            cells.append(c)

                    if time_col >= len(cells) or dlg_col >= len(cells):
                        continue

                    tc_cell  = cells[time_col]
                    dlg_cell = cells[dlg_col]

                    # Strip visual-description lines that may have leaked into the
                    # dialogue cell (e.g. "SCENE 5\nActuality:" on a shared row)
                    dlg_lines = [
                        ln.strip() for ln in dlg_cell.split("\n")
                        if ln.strip() and not _VISUAL_KEYWORDS.match(ln.strip())
                    ]

                    row_tcs = [t.strip() for t in tc_cell.split("\n")
                               if tc_re.match(t.strip())]

                    if len(row_tcs) >= 2 and dlg_lines:
                        # Multiple timecodes + dialogue in ONE row — distribute lines evenly
                        _flush_cue()
                        cur_tc[0]  = None
                        cur_dlg[0] = []
                        n = len(row_tcs)
                        grouped = {tc: [] for tc in row_tcs}
                        for di, d in enumerate(dlg_lines):
                            gi = (round(di * (n - 1) / max(1, len(dlg_lines) - 1))
                                  if len(dlg_lines) > 1 else 0)
                            gi = max(0, min(n - 1, gi))
                            grouped[row_tcs[gi]].append(d)
                        for tc in row_tcs:
                            if grouped[tc]:
                                sections.append(f"{tc} | {chr(10).join(grouped[tc])}")
                        continue

                    if len(row_tcs) == 1:
                        _flush_cue()
                        cur_tc[0]  = row_tcs[0]
                        cur_dlg[0] = list(dlg_lines)
                    elif dlg_lines:
                        cur_dlg[0].extend(dlg_lines)

                _flush_cue()

            else:
                # Unknown table layout → emit non-visual cells only
                for row in table.rows:
                    cells = []
                    seen = set()
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t in seen:
                            continue
                        if t:
                            seen.add(t)
                        # Drop cells that are purely visual/actuality descriptions
                        if t and _cell_is_visual(t):
                            continue
                        cells.append(t)
                    if any(cells):
                        sections.append(" | ".join(cells))
        extracted_text = "\n".join(sections)
    except Exception as e:
        # Fallback for poorly formed DOCX files (like missing app.xml)
        import zipfile
        import xml.etree.ElementTree as ET
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                xml_content = z.read('word/document.xml')
                root = ET.fromstring(xml_content)
                lines = []
                # Find all body elements: paragraphs and tables
                body = None
                for elem in root.iter():
                    if elem.tag.endswith('}body'):
                        body = elem
                        break
                if not body:
                    body = root

                for child in body:
                    if child.tag.endswith('}p'):
                        p_text = []
                        for t_elem in child.iter():
                            if t_elem.tag.endswith('}t') and t_elem.text:
                                p_text.append(t_elem.text)
                        joined = "".join(p_text).strip()
                        if joined:
                            lines.append(joined)
                    elif child.tag.endswith('}tbl'):
                        lines.append("=== TABLE ===")
                        for row in child.iter():
                            if row.tag.endswith('}tr'):
                                cells = []
                                for tc in row.iter():
                                    if tc.tag.endswith('}tc'):
                                        tc_text = []
                                        for t_elem in tc.iter():
                                            if t_elem.tag.endswith('}t') and t_elem.text:
                                                tc_text.append(t_elem.text)
                                        cells.append("".join(tc_text).strip())
                                if any(cells):
                                    lines.append(" | ".join(cells))
                extracted_text = "\n".join(lines)
        except:
            raise e

    # NOTE: OCR is now handled by the universal pass in read_file()
    # (_run_universal_ocr) so every embedded image is covered exactly once.

    return extracted_text


def read_legacy_doc(file_bytes: bytes) -> str:
    import tempfile
    import os
    try:
        from spire.doc import Document
    except ImportError:
        return decode_bytes(file_bytes)
        
    fd, path = tempfile.mkstemp(suffix=".doc")
    with os.fdopen(fd, 'wb') as f:
        f.write(file_bytes)
    
    try:
        from spire.doc import Document, DocumentObjectType
        document = Document()
        document.LoadFromFile(path)
        
        lines = []
        for i in range(document.Sections.Count):
            section = document.Sections.get_Item(i)
            for j in range(section.Body.ChildObjects.Count):
                obj = section.Body.ChildObjects.get_Item(j)
                if obj.DocumentObjectType == DocumentObjectType.Paragraph:
                    text = obj.Text.strip()
                    if text and "Evaluation Warning:" not in text:
                        lines.append(text)
                elif obj.DocumentObjectType == DocumentObjectType.Table:
                    table = obj
                    for r in range(table.Rows.Count):
                        row = table.Rows.get_Item(r)
                        cells = []
                        for c in range(row.Cells.Count):
                            cell = row.Cells.get_Item(c)
                            cell_text = []
                            for k in range(cell.Paragraphs.Count):
                                pt = cell.Paragraphs.get_Item(k).Text.strip()
                                if pt: cell_text.append(pt)
                            cells.append(" ".join(cell_text).strip())
                        if any(cells):
                            lines.append(" | ".join(cells))
                            
        result = "\n".join(lines).strip()
        import sys
        print("DEBUG read_legacy_doc: SUCCESS, len=", len(result), "tables_found=", len([l for l in lines if '|' in l]), file=sys.stderr)
        return result
    except Exception as e:
        import sys
        print(f"Spire.Doc error: {e}", file=sys.stderr)
        return decode_bytes(file_bytes)
    finally:
        try:
            os.remove(path)
        except:
            pass


def read_pdf(file_bytes: bytes) -> str:
    sections = []
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype='pdf')
        for i, page in enumerate(doc):
            sections.append(f'=== PAGE {i+1} ===')
            text = page.get_text('text')
            sections.append(text.strip())
        return '\n'.join(sections)
    except Exception as e:
        import sys
        print(f'PyMuPDF error: {e}', file=sys.stderr)
        try:
            import PyPDF2
            import io
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pages = []
            for i, p in enumerate(reader.pages):
                t = p.extract_text() or ''
                if t.strip():
                    pages.append(f'=== PAGE {i+1} ===\n{t}')
            return '\n'.join(pages)
        except Exception as e2:
            import sys
            print(f'PyPDF2 fallback error: {e2}', file=sys.stderr)
            return ''

def _read_pdf_spatial(file_bytes: bytes) -> str:
    """
    Spatial word-by-word PDF parser for CCSL spotting list PDFs.

    When pdfplumber's table extractor yields only header rows and no data,
    this function reconstructs the table by:
    1. Extracting all words with bounding boxes and font info.
    2. Grouping words into visual rows by y-coordinate (±3px).
    3. Detecting the header row to learn column x-boundaries.
    4. Assigning each subsequent word to a column by x-position.
    5. Returning pipe-separated rows.

    Bold/italic font names are preserved as **word** / *word* markers so the
    downstream cleaner can apply <i>/<b> tags correctly.
    The CCSL columns are:
      Sh# | ShTimeIn | SceneDescription | Title | TimeIn | TimeOut | Dur | Titles
    We care most about TimeIn, TimeOut, and Titles (the subtitle text).
    """
    import sys
    try:
        import pdfplumber
    except ImportError:
        return ""

    all_lines = []
    header_printed = False

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            # Detect global column boundaries from first page that has a header
            global_col_bounds: list[float] = []

            for page_num, page in enumerate(pdf.pages):
                try:
                    words = page.extract_words(
                        keep_blank_chars=False,
                        extra_attrs=["fontname", "size"],
                        x_tolerance=2,
                        y_tolerance=3,
                    )
                except Exception:
                    words = page.extract_words(keep_blank_chars=False)

                if not words:
                    continue

                # Group words into visual rows (snap y to 3px grid)
                rows: dict[int, list] = {}
                for w in words:
                    y_key = round(w["top"] / 3) * 3
                    if y_key not in rows:
                        rows[y_key] = []
                    text = w["text"]
                    fname = w.get("fontname", "")
                    if re.search(r'Bold|Heavy|Black', fname, re.IGNORECASE):
                        text = f"**{text}**"
                    elif re.search(r'Italic|Oblique', fname, re.IGNORECASE):
                        text = f"*{text}*"
                    rows[y_key].append((w["x0"], w["x1"], text))

                sorted_y = sorted(rows.keys())

                # Find header row on this page
                col_bounds: list[float] = list(global_col_bounds)
                header_y = None

                for y in sorted_y:
                    row_words = sorted(rows[y], key=lambda x: x[0])
                    row_text = " ".join(w[2] for w in row_words)
                    if re.search(
                    r'(Sh#|ShTimeIn|SceneDescription|Time\s*In|Time\s*Out|Titles'
                    r'|SC#|FOOTAGE|SHOT\s*DESCR|START\s*:|FINISH\s*:|TOTAL\s*:'
                    r'|TITLE#|TITLE\s*:)',
                    row_text, re.IGNORECASE
                ):
                        col_bounds = [w[0] for w in row_words]
                        if not global_col_bounds:
                            global_col_bounds = col_bounds
                        header_y = y
                        if not header_printed:
                            # Output a clean header line for downstream parsers
                            header_line = " | ".join(w[2] for w in row_words)
                            all_lines.append(header_line)
                            header_printed = True
                        break

                def assign_col(x0: float, bounds: list[float]) -> int:
                    for idx in range(len(bounds) - 1, -1, -1):
                        if x0 >= bounds[idx] - 8:
                            return idx
                    return 0

                for y in sorted_y:
                    if header_y is not None and y <= header_y:
                        continue
                    row_words = sorted(rows[y], key=lambda x: x[0])
                    if not row_words:
                        continue

                    if col_bounds and len(col_bounds) >= 4:
                        col_groups: dict[int, list[str]] = {}
                        for (x0, x1, text) in row_words:
                            col = assign_col(x0, col_bounds)
                            col_groups.setdefault(col, []).append(text)
                        num_cols = len(col_bounds)
                        cells = [" ".join(col_groups.get(c, [])).strip() for c in range(num_cols)]
                        if any(c for c in cells):
                            all_lines.append(" | ".join(cells))
                    else:
                        line = " ".join(w[2] for w in row_words).strip()
                        if line:
                            all_lines.append(line)

    except Exception as e:
        import sys
        print(f"[_read_pdf_spatial] error: {e}", file=sys.stderr)

    return "\n".join(all_lines)


def read_srt(file_bytes: bytes) -> str:
    return decode_bytes(file_bytes)


def read_vtt(file_bytes: bytes) -> str:
    return decode_bytes(file_bytes)


def read_xml(file_bytes: bytes) -> str:
    import xml.etree.ElementTree as ET
    raw = decode_bytes(file_bytes)
    lines = []
    try:
        root = ET.fromstring(raw)
        def strip_ns(tag):
            return tag.split("}")[-1] if "}" in tag else tag
        for elem in root.iter():
            tag = strip_ns(elem.tag)
            if tag == "p":
                # itertext() collects text from the <p> AND all nested spans,
                # so dialogues are never lost when text lives in child elements.
                text = " ".join(" ".join(elem.itertext()).split()).strip()
                if not text:
                    continue
                begin = end = ""
                for k, v in elem.attrib.items():
                    nk = strip_ns(k)
                    if nk == "begin":
                        begin = v
                    elif nk == "end":
                        end = v
                lines.append(f"{begin} --> {end} | {text}" if begin else text)
    except Exception:
        # Fallback: regex-strip all tag contents, keep everything.
        lines = [t.strip() for t in re.findall(r">([^<]+)<", raw) if t.strip()]
    return "\n".join(lines)


def read_rtf(file_bytes: bytes) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
        extracted_text = rtf_to_text(decode_bytes(file_bytes))
    except:
        text = decode_bytes(file_bytes)
        text = re.sub(r'\\[a-zA-Z]+\-?\d*\s?', ' ', text)
        text = re.sub(r'[{}\\]', '', text)
        extracted_text = re.sub(r'\s+', ' ', text).strip()

    # OCR fallback — RTF can embed images via \pict blocks (a pasted
    # screenshot of a table, same category of issue as DOCX/XLSX above,
    # NOTE: OCR is now handled by the universal pass in read_file()
    # (_run_universal_ocr) so every embedded image is covered exactly once.

    return extracted_text



def list_excel_sheets(file_bytes: bytes) -> list[dict]:
    """
    Return a lightweight list of sheets in an Excel file WITHOUT reading all
    cell content. Each entry: {"name": str, "row_count": int}.
    Used by the /platforms/preview-excel endpoint so the UI can show the
    subtitler which sheets exist before they decide what to import.
    """
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        sheets = []
        for ws in wb.worksheets:
            try:
                rows = ws.max_row or 0
            except Exception:
                rows = 0
            sheets.append({"name": ws.title, "row_count": rows})
        wb.close()
        return sheets
    except Exception:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_bytes)
            return [{"name": s.name, "row_count": s.nrows} for s in wb.sheets()]
        except Exception:
            return []


def read_excel_sheet(file_bytes: bytes, sheet_name: str) -> str:
    """
    Extract text from ONE named sheet of an Excel file, ignoring all other
    sheets completely. This prevents rule bleed-over when a multi-platform
    guidelines workbook is uploaded for a single OTT.

    Returns the sheet content as pipe-delimited rows (same format as
    read_excel, but for a single sheet only).
    """
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        # Case-insensitive sheet name matching
        sheet = next(
            (ws for ws in wb.worksheets if ws.title.strip().lower() == sheet_name.strip().lower()),
            None
        )
        if sheet is None:
            # Fall back to reading all sheets if the name isn't found
            return read_excel(file_bytes)
        sections = [f"=== SHEET: {sheet.title} ==="]
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                sections.append(" | ".join(cells))
                
        extracted_text = "\n".join(sections)
        # NOTE: OCR handled by universal pass in read_file() (_run_universal_ocr).
        return extracted_text
    except Exception:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_bytes)
            sheet = next(
                (s for s in wb.sheets() if s.name.strip().lower() == sheet_name.strip().lower()),
                None
            )
            if sheet is None:
                return read_excel(file_bytes)
            sections = [f"=== SHEET: {sheet.name} ==="]
            for r in range(sheet.nrows):
                cells = [str(sheet.cell_value(r, c)).strip()
                         for c in range(sheet.ncols)
                         if str(sheet.cell_value(r, c)).strip()]
                if cells:
                    sections.append(" | ".join(cells))
            
            extracted_text = "\n".join(sections)
            # NOTE: OCR handled by universal pass in read_file() (_run_universal_ocr).
            return extracted_text
        except Exception:
            return ""


def read_excel(file_bytes: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        sections = []
        for sheet in wb.worksheets:
            sections.append(f"=== SHEET: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    sections.append(" | ".join(cells))
        extracted_text = "\n".join(sections)

        # NOTE: OCR handled by universal pass in read_file() (_run_universal_ocr).
        return extracted_text
    except:
        # Legacy .xls binary format — NOT a zip archive, so the OCR path
        # above (which relies on zipfile) does not apply here. Falls back
        # to plain text extraction only.
        import xlrd
        wb = xlrd.open_workbook(file_contents=file_bytes)
        sections = []
        for sheet in wb.sheets():
            sections.append(f"=== SHEET: {sheet.name} ===")
            for r in range(sheet.nrows):
                cells = [str(sheet.cell_value(r, c)).strip()
                         for c in range(sheet.ncols)
                         if str(sheet.cell_value(r, c)).strip()]
                if cells:
                    sections.append(" | ".join(cells))
        extracted_text = "\n".join(sections)

        # NOTE: OCR handled by universal pass in read_file() (_run_universal_ocr).
        return extracted_text


def read_csv(file_bytes: bytes) -> str:
    import csv
    text = decode_bytes(file_bytes)
    reader = csv.reader(io.StringIO(text))
    lines = [" | ".join(c.strip() for c in row if c.strip()) for row in reader]
    return "\n".join(l for l in lines if l)


def read_plain(file_bytes: bytes) -> str:
    return decode_bytes(file_bytes)


def read_json(file_bytes: bytes) -> str:
    import json
    try:
        data = json.loads(decode_bytes(file_bytes))
        return json.dumps(data, indent=2, ensure_ascii=False)
    except:
        return decode_bytes(file_bytes)
