# extractor.py — Pure Python dialogue extraction (zero LLM)
# Based on real company cleaning process from How_to_Clean_Script_Updated_.pdf
# Handles COYOTE-style SUBTITLE rows, GRAPHICS ON SCREEN removal, all bracket types

import re

# ── Patterns ────────────────────────────────────────────────────────────────────

_TIMECODE = re.compile(
    r'^\d{1,2}:\d{2}:\d{2}[;:,.]\d{2,3}\s*'       # leading timecode HH:MM:SS:FF
    r'|\d{1,2}:\d{2}:\d{2}\s*-->\s*\d{1,2}:\d{2}:\d{2}'  # SRT arrow
)

_SCENE_HEADINGS = re.compile(
    r'^(INT\.|EXT\.|INT/EXT\.|EXT/INT\.|TEASER|TAG\b|ACT\s+(ONE|TWO|THREE|\d+)'
    r'|SCENE\s+\d+|COLD\s+OPEN|FADE\s+IN|FADE\s+OUT|CUT\s+TO|SMASH\s+CUT)',
    re.IGNORECASE
)

# ALL-CAPS words that are production notes, not dialogue
_ALLCAPS_NOISE = re.compile(
    r'^(WALLA|MUSIC|CHATTER'
    r'|INTERCUT|FLASHBACK|MONTAGE|SUPER)[:\s]',
    re.IGNORECASE
)

# Lines that MUST be preserved — they're needed by platform rules
# e.g. "END CREDITS", "TITLE SEQUENCE", on-screen text, opening/closing titles
_MUST_PRESERVE = re.compile(
    r'(END\s+CREDIT|OPENING\s+CREDIT|CLOSING\s+CREDIT|TITLE\s+CARD|TITLE\s+SEQUENCE'
    r'|ON\s+SCREEN|ON-SCREEN|ONSCREEN|LOWER\s+THIRD'
    r'|GRAPHIC|SUPER(?:IMPOSE)?\b|CAPTION\b|SUBTITLE\b'
    r'|ROLL\s+CREDIT|CREDIT\s+ROLL|CREDITS\s+ROLL'
    r'|WRITTEN\s+BY|DIRECTED\s+BY|PRODUCED\s+BY'
    r'|EXECUTIVE\s+PRODUCER|STARRING|FEATURING'
    r'|TRANSLATION|SUBTITLING\s+BY|TRANSLATED\s+BY)',
    re.IGNORECASE
)


def _is_credits_or_preserved(line: str) -> bool:
    """Return True if this line is credit/title/on-screen content that must not be filtered."""
    return bool(_MUST_PRESERVE.search(line))

# Stage directions in ANY bracket type — catches ALL parenthetical content
_BRACKETS = re.compile(
    r'\(speaking\s+\w+\)\s*'           # (speaking Spanish)
    r'|\(through\s+\w+\)\s*'           # (through telephone)
    r'|\(into\s+\w+\)\s*'              # (into phone) (into radio)
    r'|\(to\s+[^)]+\)\s*'              # (to himself) (to Maria)
    r'|\(whispering[^)]*\)\s*'         # (whispering in Spanish)
    r'|\(overlaps?\)\s*'               # (overlaps)
    r'|\(Archive\)\s*'                 # (Archive)
    r'|\(continues?[^)]*\)\s*'         # (continues...)
    r'|\(indistinct[^)]*\)\s*'         # (indistinct chatter)
    r'|\(.*?music.*?\)\s*'             # (music)
    r'|\(sarcastically\)\s*'           # (sarcastically)
    r'|\(quietly\)\s*'                 # (quietly)
    r'|\(loudly\)\s*'                  # (loudly)
    r'|\(angrily\)\s*'                 # (angrily)
    r'|\(softly\)\s*'                  # (softly)
    r'|\(nervously\)\s*'               # (nervously)
    r'|\(in\s+\w+\)\s*'               # (in English) (in Spanish)
    r'|\[.*?\]'                        # [gasps] [MUSIC] [indistinct]
    r'|<[^>]+>'                        # <gasps> <laughs>
    r'|\(LAUGHS?\)|\(CHUCKLES?\)|\(GASPS?\)|\(SIGHS?\)|\(SCREAMS?\)'
    r'|\(CRIES?\)|\(SOBBING\)|\(MOANS?\)|\(GROANS?\)'
    r'|\([^)]+:\s*[^)]+\)',            # Slang notes like (Come on: interjection...)
    re.IGNORECASE
)

# Character name patterns — speaker labels (ends with : or -)
_SPEAKER_LABEL = re.compile(
    r'^(?:\(OPTIONAL\)\s*)?[A-Z][A-Z\s\.\-\'\/\(\)0-9\#\&\,]{1,60}[:\-]\s*'
)

# Inline speaker without colon/hyphen (ALL CAPS followed by space)
_INLINE_SPEAKER = re.compile(
    r"^([A-Z]{2,30}(?:\s+[A-Z0-9]{1,30}){0,3}|[A-Z]\s+[A-Z0-9]{1,30})\s+(?=[A-Z][a-z]|I\s|['\"\(]|\d)"
)

# Lines to skip entirely if they contain these metadata phrases (used with search)
_SKIP_PATTERNS = re.compile(
    r'(WALLA\b|MUSIC\b'
    r'|INTERCUT|FLASHBACK|MONTAGE'
    r'|DIALOGUE\s+LIST|SPOTTING\s+LIST|TIMECODE\/ DIALOGUE|START\s+MEASURING'
    r'|LOCATION\s+START:|TITLE\#:|FINISH:|TOTAL:)',
    re.IGNORECASE
)

# Comma cleanup — remove comma before connectors (company's step 5)
_COMMA_CONNECTORS = re.compile(r',\s+(and|but|that|or|because)\b', re.IGNORECASE)


def _clean_line(line: str) -> str:
    """Strip stage-direction noise from a single line — preserves <i>/<b> italic/bold tags."""
    # Preserve italic/bold tags
    line = re.sub(r'<i>', '__ITALIC_OPEN__', line)
    line = re.sub(r'</i>', '__ITALIC_CLOSE__', line)
    line = re.sub(r'<b>', '__BOLD_OPEN__', line)
    line = re.sub(r'</b>', '__BOLD_CLOSE__', line)
    # Remove bracket-based stage directions
    line = _BRACKETS.sub('', line)
    # Remove other markup tags (but not our placeholders)
    line = re.sub(r'<[^>]+>', '', line)
    # Fix multiple spaces
    line = re.sub(r'\s+', ' ', line).strip()
    # Remove comma before connectors
    line = _COMMA_CONNECTORS.sub(r' \1', line)
    # Strip leading/trailing punctuation artefacts
    line = line.strip(' -\u2013\u2014|')
    # Restore italic/bold tags
    line = line.replace('__ITALIC_OPEN__', '<i>')
    line = line.replace('__ITALIC_CLOSE__', '</i>')
    line = line.replace('__BOLD_OPEN__', '<b>')
    line = line.replace('__BOLD_CLOSE__', '</b>')
    return line


def _is_valid(line: str, min_words: int = 1) -> bool:
    """Check if line is real dialogue worth keeping.
    Credit/title/on-screen lines bypass all filters — they are needed for platform rules.
    """
    line = line.strip()
    if not line or len(line) < 2:
        return False
    # Credits and platform-rule-required content must ALWAYS pass through
    if _is_credits_or_preserved(line):
        return True
    if _SCENE_HEADINGS.match(line):
        return False
    if _SKIP_PATTERNS.search(line):
        return False
    # Check plain-text length (strip tags for this check)
    plain = re.sub(r'<[^>]+>', '', line)
    if _TIMECODE.match(plain) and len(plain.replace(' ', '')) < 20:
        return False
    if re.match(r'^[\d\s\.\,\-\:\;\/]+$', plain):
        return False
    if plain.startswith('(') and plain.endswith(')'):
        return False
    if plain.startswith('[') and plain.endswith(']'):
        return False
    words = [w for w in plain.split() if re.search(r'[a-zA-Z\u00C0-\u024F]', w)]
    return len(words) >= min_words


# ── DOCX TABLE EXTRACTOR ────────────────────────────────────────────────────────

def extract_docx_table(file_bytes: bytes) -> list[str]:
    """
    Smart DOCX table extractor.
    Handles:
    - Standard tables: TIMECODE | CHARACTER | DIALOGUE
    - SUBTITLE rows (keep — these are the English translations)
    - NARRATIVE TITLE, GRAPHICS ON SCREEN rows (skip)
    - Stage directions in parentheses within dialogue (strip)
    - <gasps>, [MUSIC] etc (strip)
    Based on COYOTE_102 real file structure.
    """
    from docx import Document
    import io
    doc = Document(io.BytesIO(file_bytes))
    lines = []

    for table in doc.tables:
        cols = len(table.columns)
        for row in table.rows:
            cells = []
            seen = set()
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in seen:
                    cells.append(t)
                    seen.add(t)

            if not cells:
                continue

            # Skip header row
            if cells[0].upper() in ('TIME CODE', 'TIMECODE', 'TC'):
                continue

            # For 3-column or more tables: TIMECODE | CHARACTER | DIALOGUE, or Spotting Lists
            if len(cells) >= 3:
                # Look for the last valid dialogue cell, but skip the 'TOTAL' or timecode columns usually at the end
                dialogue = ''
                for cell in reversed(cells):
                    t_clean = _clean_line(cell)
                    if _is_valid(t_clean) and not _TIMECODE.match(cell) and not re.match(r'^[\d\.]+$', cell.strip()):
                        dialogue = cell
                        break

                # Skip rows with no dialogue
                if not dialogue or not dialogue.strip():
                    continue

                # Skip GRAPHICS ON SCREEN, NARRATIVE TITLE, WALLA etc
                # BUT: preserve credit/title/on-screen lines needed for platform rules
                dlg_stripped = dialogue.strip()
                if not _is_credits_or_preserved(dlg_stripped):
                    if _SKIP_PATTERNS.match(dlg_stripped):
                        continue
                    if _ALLCAPS_NOISE.match(dlg_stripped):
                        continue

                # Clean and strip speaker labels / slang notes
                cleaned_dialogue = _clean_line(dialogue)
                cleaned_dialogue = _SPEAKER_LABEL.sub('', cleaned_dialogue).strip()
                m = _INLINE_SPEAKER.match(cleaned_dialogue)
                if m: cleaned_dialogue = cleaned_dialogue[m.end():].strip()
                cleaned_dialogue = re.sub(r'\([^)]+\)$', '', cleaned_dialogue).strip()
                
                if not cleaned_dialogue:
                    continue

                # Skip sound-only rows like <gasps>, [MUSIC]
                if not _is_valid(cleaned_dialogue):
                    continue

                lines.append(cleaned_dialogue)

            # For 2-column or unusual tables
            elif len(cells) >= 1:
                dialogue = cells[-1]
                if _SKIP_PATTERNS.match(dialogue.strip()):
                    continue
                cleaned = _clean_line(dialogue)
                cleaned = _SPEAKER_LABEL.sub('', cleaned).strip()
                m = _INLINE_SPEAKER.match(cleaned)
                if m: cleaned = cleaned[m.end():].strip()
                cleaned = re.sub(r'\([^)]+\)$', '', cleaned).strip()
                if _is_valid(cleaned):
                    lines.append(cleaned)

    # Also check paragraphs (some docs have content outside tables)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Strip timecodes from paragraph text
        text = _TIMECODE.sub('', text).strip()
        if _SCENE_HEADINGS.match(text) or _SKIP_PATTERNS.match(text):
            continue
        text = _SPEAKER_LABEL.sub('', text)
        m = _INLINE_SPEAKER.match(text)
        if m: text = text[m.end():].strip()
        cleaned = _clean_line(text)
        cleaned = re.sub(r'\([^)]+\)$', '', cleaned).strip()
        if _is_valid(cleaned, min_words=2):
            lines.append(cleaned)

    return _dedup(lines)


# ── SRT EXTRACTOR ────────────────────────────────────────────────────────────────

def extract_srt(text: str) -> list[str]:
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line or line.isdigit() or '-->' in line:
            continue
        cleaned = _clean_line(line)
        if _is_valid(cleaned):
            lines.append(cleaned)
    return _dedup(lines)


# ── VTT EXTRACTOR ────────────────────────────────────────────────────────────────

def extract_vtt(text: str) -> list[str]:
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line or '-->' in line or line.startswith('WEBVTT') or line.startswith('NOTE'):
            continue
        if line.isdigit():
            continue
        cleaned = _clean_line(line)
        if _is_valid(cleaned):
            lines.append(cleaned)
    return _dedup(lines)


# ── TABLE TEXT EXTRACTOR (pipe-separated from file_reader) ───────────────────────

def extract_table(text: str) -> list[str]:
    """For pipe-separated table text already processed by file_reader"""
    lines = []
    for raw_line in text.splitlines():
        if '|' not in raw_line:
            continue
        if raw_line.strip().startswith('==='):
            continue
        cells = [c.strip() for c in raw_line.split('|')]

        # Try to find dialogue in last meaningful cell
        dialogue = ''
        valid_cells = []
        for i, cell in enumerate(cells):
            cell_c = _clean_line(cell)
            # Skip purely numeric cells (like durations) or timecodes
            if _is_valid(cell_c) and not _TIMECODE.match(cell) and not re.match(r'^[\d\.\:]+$', cell.strip()):
                valid_cells.append((i, cell_c))

        if valid_cells:
            # If it's a double dialogue spotting list and we have multiple valid cells, 
            # the last valid cell is typically the Subtitle column. We shouldn't fall back 
            # to earlier cells (like the Transcript column) if this row genuinely lacks a subtitle.
            # But since we iterate row by row, we just take the rightmost valid cell.
            dialogue = valid_cells[-1][1]

            # If we picked the first column but there were >=3 columns total, it's probably 
            # just transcript audio without a subtitle. Skip it to avoid noise!
            if valid_cells[-1][0] == 0 and len(cells) >= 3 and re.search(r'\(OS|\(VO|\(CONT|\(OVERLAPPING', cells[0], re.IGNORECASE):
                dialogue = ''

        if dialogue and not _SKIP_PATTERNS.search(dialogue):
            # Remove speaker label if present in the dialogue cell
            dialogue = _SPEAKER_LABEL.sub('', dialogue).strip()
            m = _INLINE_SPEAKER.match(dialogue)
            if m: dialogue = dialogue[m.end():].strip()
            # Remove trailing slang definitions
            dialogue = re.sub(r'\([^)]+\)$', '', dialogue).strip()
            if _is_valid(dialogue):
                lines.append(dialogue)

    return _dedup(lines)


# ── PARAGRAPH WITH SPEAKER ───────────────────────────────────────────────────────

def extract_script_with_speaker(text: str) -> list[str]:
    """
    Handles: TIMECODE  CHARACTER NAME  dialogue text
    Or: CHARACTER NAME:\ndialogue
    Based on the exact format shown in How_to_Clean_Script PDF
    """
    lines = []
    
    # Identify if we are in a spotting list/CCSL where we should filter out transcript-only lines
    is_ccsl = "COMBINED CONTINUITY" in text.upper() or "SPOTTING LIST" in text.upper() or "CCSL" in text.upper()
    
    # Transcript-specific speaker tags to ignore in spotting lists
    transcript_tags = re.compile(r'\(OS\)|\(VO\)|\(CONT\)|\(OVERLAPPING\)|\(OS CONT\)|\(OS OVERLAPPING\)', re.IGNORECASE)

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        # Remove leading timecode
        line = _TIMECODE.sub('', line).strip()
        if not line:
            continue
        # Skip scene headings and noise — but always preserve credits/on-screen text
        if _is_credits_or_preserved(line):
            # Let credits through without any further filtering
            lines.append(_clean_line(line))
            continue
        if _SCENE_HEADINGS.match(line) or _SKIP_PATTERNS.search(line):
            continue
        # Remove ALL CAPS lines that are scene/production notes
        # EXCEPTION: do NOT filter if it looks like credits, titles, or on-screen text
        if re.match(r'^[A-Z0-9\s\.\-\'\/\(\)\&\,]{3,}$', line):
            continue
            
        # In flattened spotting lists, skip transcript lines (they have OS, VO, etc.)
        if is_ccsl and transcript_tags.search(line):
            continue
            
        # Optional: remove subtitle numbers like '4-1 ' at the start of subtitles
        line = re.sub(r'^\d+\-\d+\s+', '', line)
            
        # Remove speaker label
        line = _SPEAKER_LABEL.sub('', line).strip()
        m = _INLINE_SPEAKER.match(line)
        if m: line = line[m.end():].strip()
        # Also remove any remaining slang definitions like (Slang)
        line = re.sub(r'\([^)]+\)$', '', line).strip() # Removes trailing brackets like (Be quiet)
        cleaned = _clean_line(line)
        if _is_valid(cleaned, min_words=1):
            lines.append(cleaned)
    return _dedup(lines)


# ── PARAGRAPH WITHOUT TABLE (timecodes inline) ───────────────────────────────────

def extract_script_with_timecodes(text: str) -> list[str]:
    lines = []
    for line in text.splitlines():
        line = line.strip()
        line = _TIMECODE.sub('', line).strip()
        if not line:
            continue
        if _SCENE_HEADINGS.match(line) or _SKIP_PATTERNS.search(line):
            continue
        line = _SPEAKER_LABEL.sub('', line).strip()
        m = _INLINE_SPEAKER.match(line)
        if m: line = line[m.end():].strip()
        cleaned = _clean_line(line)
        cleaned = re.sub(r'\([^)]+\)$', '', cleaned).strip()
        if _is_valid(cleaned, min_words=1):
            lines.append(cleaned)
    return _dedup(lines)


# ── PLAIN SCRIPT ─────────────────────────────────────────────────────────────────

def extract_plain(text: str) -> list[str]:
    """For already-cleaned RTF/TXT scripts — minimal processing"""
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        cleaned = _clean_line(line)
        if _is_valid(cleaned, min_words=1):
            lines.append(cleaned)
    return _dedup(lines)


# ── DEDUP ────────────────────────────────────────────────────────────────────────

def _dedup(lines: list[str]) -> list[str]:
    """Remove consecutive duplicate lines"""
    result = []
    prev = None
    for line in lines:
        if line != prev:
            result.append(line)
            prev = line
    return result


# ── DISPATCHER ───────────────────────────────────────────────────────────────────

def pre_extract_dialogue(raw_text: str, structure: str, file_bytes: bytes = None, filename: str = '', platform_dict: dict = None) -> list[str]:
    """
    Main entry point.
    For DOCX files with tables — uses smart DOCX table extractor directly.
    For all other formats — uses text-based extractors.
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if platform_dict is None:
        platform_dict = {}

    # DOCX gets special treatment — read table cells directly
    if ext in ('docx', 'doc') and file_bytes:
        try:
            result = extract_docx_table(file_bytes)
            if result:
                print(f"[EXTRACT] DOCX table: {len(result)} dialogue lines extracted")
                return result
        except Exception as e:
            print(f"[EXTRACT] DOCX table failed: {e}, falling back to text")

    # All other formats use text-based extraction
    extractors = {
        'srt_format':              extract_srt,
        'vtt_format':              extract_vtt,
        'table_with_timecodes':    extract_table,
        'ccsl_double_dialogue':    extract_table,
        'excel_spotting_list':     extract_table,
        'paragraph_with_speaker':  extract_script_with_speaker,
        'paragraph_without_table': extract_script_with_timecodes,
        'plain_script':            extract_plain,
        'xml_ttml':                extract_srt,
        'unknown':                 extract_script_with_speaker,
    }

    fn = extractors.get(structure, extract_plain)
    result = fn(raw_text)
    
    # Fallback if structure detection was wrong and it extracted almost nothing
    if len(result) < 10 and len(raw_text.splitlines()) > 30 and fn != extract_script_with_speaker:
        print(f"[EXTRACT] {structure} failed (only {len(result)} lines). Falling back to extract_script_with_speaker.")
        result = extract_script_with_speaker(raw_text)
        
    # Post-process based on platform guidelines
    remove_elements = platform_dict.get("remove_elements", [])
    
    final_result = []
    for line in result:
        # If fillers should be removed
        if "fillers" in remove_elements:
            line = re.sub(r'\b(ugh|hmm|erm|ah|oh)\b', '', line, flags=re.IGNORECASE)
            line = re.sub(r'\s+', ' ', line).strip()
            
        # Optional: remove purely uppercase lines if they are not dialogue but scene descriptions
        # CRITICAL: never remove lines that contain credit/platform-rule-required content
        if "scene_descriptions" in remove_elements and line.isupper() and len(line) > 5:
            if _is_credits_or_preserved(line):
                pass  # Always keep credits regardless of scene_descriptions filter
            elif re.match(r'^[A-Z0-9\s\.\-\'\/\(\)\.\&\,]{3,60}$', line) and not any(p in line for p in ['?', '!', '"']):
                continue
                
        if not line or len(line) < 2:
            continue
            
        final_result.append(line)
        
    result = final_result

    print(f"[EXTRACT] {structure}: {len(result)} lines via Python extractor")
    return result
